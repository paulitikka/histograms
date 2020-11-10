# -*- coding: utf-8 -*-
"""
Created on Mon Jul 20 10:46:49 2020 (Latest update from: 5.11.20 by Pauli Tikka)

@author: pauli
"""
#Tikka et al 2020 Peer Review Word Count Histograms etc. codes (4.11.20, Tikka)
#%% First import all packages:
import requests
import urllib.request
import time
from bs4 import BeautifulSoup
import re
import pandas as pd #for importing files
# https://pandas.pydata.org/pandas-docs/version/0.18.1/generated/pandas.DataFrame.html
import numpy as np  #for calculations, array manipulations, and fun :)
import matplotlib.pyplot as plt #for scientifical plots
import os
import seaborn
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import datetime
import time
from selenium import webdriver  # for webdriver
from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.action_chains import ActionChains
from selenium.webdriver.support.ui import WebDriverWait  # for implicit and explict waits
from selenium.webdriver.chrome.options import Options  # for suppressing the browser
import requests
from bs4 import BeautifulSoup
from concurrent.futures import ProcessPoolExecutor, as_completed
from selenium.webdriver.common.by import By
import random
import string
from matplotlib.legend import Legend
import matplotlib.lines as mlines   
import datetime
import numpy as np
import matplotlib.pyplot as plt
from matplotlib.ticker import PercentFormatter
import pandas as pd #for importing files

#%% The next codes are for Figure 1 in manuscript (ver 3.11.2020, Tikka)
#%Basic histogram for all count data first
#Bring the data in fight form:
tot_res2=pd.read_csv('all_journals_tikka21720_ok.csv', index_col=None, header=0)
#% Auxialiary, but imprtant variables for histo and bar functions:
#the dates or yearly conditions, make this in all the cases:
tot_res2.rename(columns={'Date of Publication':'Date'}, inplace=True)
#https://stackoverflow.com/questions/20868394/changing-a-specific-column-name-in-pandas-dataframe
tot_res2['Date'] =pd.to_datetime(tot_res2.Date)
#%Thrid condition, the year
#https://stackoverflow.com/questions/25146121/extracting-just-month-and-year-separately-from-pandas-datetime-column
tot_res2['year'] = pd.DatetimeIndex(tot_res2['Date']).year  
fonta=['Calibri', 20,'light',14] #To make these fonts show in graph, you need to enter twice..
fig, ax = plt.subplots()
plt.rcParams["font.family"] = fonta[0]
plt.rcParams["font.size"] = fonta[1]
plt.rcParams["font.weight"] = fonta[2]
plt.rcParams["axes.labelweight"] = fonta[2]
ax.spines['right'].set_visible(False)
ax.spines['top'].set_visible(False)
magazine=['BMC, BMJ and PLOS'] #this can be any other, e.g. 'BMJ', 'PLOS', or 'BMC, BMJ and PLOS' (i.e. ALL)
title=magazine[0] + ' in 2003-2020'   #+titleX[0][0]  #the last one could be also 'titleX' variable for any year, see below
plt.tight_layout()
rng=tot_res2['Review Word Count']
plt.hist(rng, bins=40, weights=np.ones(len(rng)) / len(rng),range=[0, 1600],alpha=1,label=title)
plt.gca().yaxis.set_major_formatter(PercentFormatter(1))
ax.legend(loc=(0.23,1),frameon=False,fontsize=fonta[3]) #You may want to specify the location: (0.35,0.97)
#https://stackoverflow.com/questions/51473993/plot-an-histogram-with-y-axis-as-percentage-using-funcformatter
plt.xlabel("Number of Words", size=fonta[1])
plt.ylabel("Review Prevalence (%)", size=fonta[1])
x_ticks = np.arange(0, 2100, 300) #the last is not 2021, but 2021-1, so for 2020 you need 2021
plt.yticks(size=fonta[1])
plt.xticks(x_ticks,size=fonta[1])
labels=[['A)']] #note the double brackets, it works like that..
blue_line = mlines.Line2D([], [], linewidth=0, marker='',markersize=0)#color='blue', marker='*',\
#        https://matplotlib.org/3.3.0/api/_as_gen/matplotlib.lines.Line2D.html
#The location of this 'A)' legend needs to be fixed:
leg = Legend(ax,labels=labels[0],loc=(-.17,1.01), handles=[blue_line],\
             handlelength=0, labelspacing=0,frameon=False,fontsize=fonta[1])
ax.add_artist(leg)
plt.tight_layout()
plt.margins(0.1, 0.1) #These are important
my_dpi=1600
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
document= Document()
document.add_heading("My results!")
font = document.styles['Normal'].font
font.name = 'Calibri'
font.size =Pt(10)
p = document.add_paragraph("Peer Review Analysis \n")
p.style = document.styles['Normal']
r = p.add_run()   
plt.savefig("Histogram of all journal years and words_tikka15920.jpg",dpi=my_dpi)
r.add_picture('C:\python\Histogram of all journal years and words_tikka15920.jpg', width = Inches(3))
plt.show()
#%At this stage, we are interested only all years or specific journals, so the very below variables 'X' are fyi
titleX=[[" in 2003-2004"],[" in 2005-2006"],[" in 2007-2008"],[" in 2009-2010"],\
        [" in 2011"],[" in 2012"],[" in 2013"],[" in 2014"],[" in 2015"],[" in 2016"],
        [" in 2017"],[" in 2018"],[" in 2019"],[" in 2020"]] 
timeX=[[2003,2004],[2005,2006],[2007,2008],[2009,2010],\
       [2011],[2012],[2013],[2014],\
       [2015],[2016],[2017],[2018],
       [2019],[2020]] 
#https://www.kite.com/python/answers/how-to-make-a-list-of-the-alphabet-in-python
import string
alphabet_string = list(string.ascii_uppercase) #fyi, needed for multiplotting..
alphabet_string2=[]
for i in range(len(alphabet_string)):
    alphabet_string2.append(alphabet_string[i]+')')
ABCs=alphabet_string2
#%Adding more contraints to the histograms, and plotting them one-by-one for each jorunal
ok=tot_res2['Review Word Count']<500 #this can be what you whant
ok_bmc=tot_res2['Journal Name']=='BMC Medicine'
ok_bmj=tot_res2['Journal Name']=='BMJ'
ok_plos=tot_res2['Journal Name']=='PLOS Medicine'
BMC=tot_res2.loc[ok_bmc]
BMJ=tot_res2.loc[ok_bmj]
PLOS=tot_res2.loc[ok_plos]
Tot=tot_res2 #%This works for all..
aza3=Tot.pivot_table(index=['Title of Article'], aggfunc='size') #this is for all
#aza3=PLOS3.pivot_table(index=['Title of Article'], aggfunc='size') #this is for individuals
aza3=pd.DataFrame(aza3)
baza3=list(aza3[0])
x=np.array(np.unique(baza3, return_counts=True))
fig, ax = plt.subplots()
fonta=['Calibri', 20,'light',14] #The same with the fonts here (as above), but should you wnat to change anything, 
#I inserted them again here:
plt.rcParams["font.family"] = fonta[0]
plt.rcParams["font.size"] = fonta[1]
plt.rcParams["font.weight"] = fonta[2]
plt.rcParams["axes.labelweight"] = fonta[2]
MEDIUM_SIZE = fonta[1]
plt.rc('xtick', labelsize=MEDIUM_SIZE) 
plt.rc('ytick', labelsize=MEDIUM_SIZE) 
ax.spines['right'].set_visible(False)
ax.spines['top'].set_visible(False)
magazine=['BMC, BMJ and PLOS'] #this can be any other, e.g. 'BMJ', 'PLOS', or 'BMC, BMJ and PLOS' (i.e. ALL)
title=magazine[0] + ' in 2003-2020'
#magazine=['PLOS'] #this can be any other, see above
#title= magazine[0]+ ' in 2019-2020'
#plt.hist(a, bins=16,range=[b,c],label=title,width=0.5)  # arguments are passed to np.histogram
ax.bar(list(x[0]), list(x[1]),align='center', width=0.6,alpha=1,label=title)
ax.legend(loc=(0.26,0.99),frameon=False,fontsize=fonta[3]) #Check the location: 0 or (0.27,0.97)/'upper right
#%Bar graph of the number of reviews per article to reach the first decision and other criteria
plt.xlabel("Reviews for First Decision", size=fonta[1])
plt.ylabel("Articles Reviewed", size=fonta[1])
y_ticks=np.arange(0, 1400, 200)
plt.yticks(y_ticks,size=fonta[1])
x_ticks=np.arange(1, np.max(x[0])+1, 1)
plt.xticks(x_ticks, size=fonta[1])
labels=[['B)']] #note the double brackets, it works like that..
blue_line = mlines.Line2D([], [], linewidth=0, marker='',markersize=0)#color='blue', marker='*',\
#https://matplotlib.org/3.3.0/api/_as_gen/matplotlib.lines.Line2D.html
#The location of this 'A)' legend needs to be fixed:
leg = Legend(ax,labels=labels[0],loc=(-.17,1.01), handles=[blue_line],\
             handlelength=0, labelspacing=0,frameon=False,fontsize=fonta[1])
ax.add_artist(leg)
plt.tight_layout()
plt.margins(0.1, 0.1) #These are important
#plt.axis((-0.5,10,-0.5,1200)) #check this if needed..
#plt.savefig("Histogram__less than 500 words_allyears_tikka4820.png")
#plt.title("All reviewed years for all journal less than 500 words")
my_dpi=1600   
plt.savefig("Histogram of all reviews_tikka18920.jpg",dpi=my_dpi)
r.add_picture('C:\python\Histogram of all reviews_tikka18920.jpg', width = Inches(3))
document.save("C:\\python\\resultna.docx")
#figsize=(800/my_dpi, 800/my_dpi)
plt.show()
 
#%%Interquortile ranges for each journal in each year, Table 1 (or S1) in manu (3.11.20, Tikka):
a = []
b = []
c = []
d = []
q25 = []
q75 = []
av1=[]
av2=np.unique(tot_res2['year'])
#%For PLOS you need revising of BMCtot and av2 values via just deleting the rows of matrices:
#av2=av2[16:] #for BMJ it starts from '12' (year 2015) or something
BMCtot=[]
for i in range(len(av2)):
    av1.append(tot_res2['year']==av2[i])
    BMCtot.append(tot_res2.loc[ok_bmc & av1[i],'Review Word Count']) 
    #check the condition 'ok_bmc' in other journals, ok_bmj/ok_plos
#    https://pandas.pydata.org/pandas-docs/stable/user_guide/merging.html
for i in range(len(av2)):
    a.append(len(BMCtot[i]))
    b.append(np.round(np.mean(BMCtot[i]),1))
    c.append(np.round(np.std(BMCtot[i]),1))
    d.append(np.round(np.median(BMCtot[i]),0))
    q25.append(np.round(np.percentile(BMCtot[i], [25]),1)[0])
    q75.append(np.round(np.percentile(BMCtot[i], [75]),1)[0])
    #%
totaal=[]
totaal=pd.concat([pd.DataFrame(a),pd.DataFrame(b),pd.DataFrame(c),\
                  pd.DataFrame(d),pd.DataFrame(q25),pd.DataFrame(q75)],axis=1)
#If you make a matrix, always add column and row (i.e. index) names:
totaal.index=av2 #the years
totaal.columns=['Reviews','Mean of Words','Std','Median','Q25','Q75'] 
#first column is the amount of reviews in that year ('Reviews'), and the rest relate to the Word Coounts
#%%Figure 2 in manuscript (ver. 3.11.2020, Tikka)
#The number of reviews (not articles) calculation: 
tot_articles_t=tot_res2[[tot_res2.columns[1],'Journal Name','year', 'Reviewer Name']]
#https://cmdlinetips.com/2019/03/how-to-select-one-or-more-columns-in-pandas/
tot_articles=len(np.unique(tot_articles_t.loc[:,tot_articles_t.columns[1]])) #'Title of Article'
#%Cheking just one year
condi=tot_articles_t.loc[:,'year']==2003
tay=tot_articles_t[condi]
year_articles=len(np.unique(tay.loc[:,'Reviewer Name']))
#%Now expanding this to all journals and years:
lena=np.unique(tot_articles_t.loc[:,'year'])
cond=[]
tayear=[]
yebye=[]
for i in range(len(lena)):
    cond.append(tot_articles_t.loc[:,'year']==lena[i])
    tayear.append(tot_articles_t[cond[i]])
    yebye.append(len(tayear[i]))
#%For BMC
conda=[]
conda2=tot_articles_t.loc[:,'Journal Name']=='PLOS Medicine'
lena=np.unique(tot_articles_t.loc[conda2,'year'])
tayeara=[]
yebyea=[]
for i in range(len(lena)):
    conda.append(tot_articles_t.loc[:,'year']==lena[i])
    tayeara.append(tot_articles_t[conda[i] & conda2])
    yebyea.append(len(tayeara[i]))    
#tot_articles03=len(np.unique(tot_res2.loc[:,'Title of Article']))
#% All reviews (no_below500_tot) or journal revious (.._bmc) that used number of words below 500:
#BMJ senior word count:
senior1=[]
senior2=[]
bmjt_t=tot_res2[tot_res2['Journal Name']=='BMJ']
#remember to reset the indeces
bmjt_t=bmjt_t.reset_index()
testSen=['Professor',  'professor', 'Prof', 'prof', 'Prof.', 'prof.', 'Dean', 'dean', 'Director', 'director', 'Head', 'head', 'Chief', 'chief', 'Chair', 'chair']
testS2=['Associate', 'associate', 'Assoc', 'assoc', 'Assoc.', 'assoc.', 'Assistant', 'assistant', 'Assist', 'assist', 'Assist.',  'assist.']
s2=[]
for i in range(len(bmjt_t)):
    for k in range(len(testSen)):
        if str(testSen[k]) in str(bmjt_t.loc[i]["Reviewer's Title"]): 
            senior1.append([i,k])
            senior2.append(bmjt_t.loc[i][bmjt_t.columns[1]])
sn2=np.unique(senior2) 
PLOS3=bmjt_t.set_index(bmjt_t.columns[1])
lehs=PLOS3.loc[sn2]
aza3=lehs.pivot_table(index=[bmjt_t.columns[1]], aggfunc='size')      
sx=[]
for i in range(len(senior1)):
    sx.append(senior1[i][0])
sn=np.unique(sx)
bok=bmjt_t.loc[sn]
bok=bok.reset_index()
st=[]
z=0
i=0
for i in range(len(bok)):
    for z in range(len(testS2)):
        if str(testS2[z]) not in str(bok.loc[i]["Reviewer's Title"]):
            st.append([i,z])
sx2=[]
for i in range(len(st)-1):
    if st[i][1]+1!=st[i+1][1]:
        sx2.append(st[i][0])
sx3=pd.DataFrame(sx2)[0]
sx4=sx3.drop_duplicates(keep=False,inplace=False) 
boka=bok.loc[sx4]
sap=[]
boka=boka.iloc[:,2:]
boka=boka.reset_index()
for i in range(len(boka)):
    sap.append(boka.loc[i]['Review Word Count'])
#%Percentage for all with one..
rng=sap
##https://stackoverflow.com/questions/12125880/changing-default-x-range-in-histogram-matplotlib 
##https://www.geeksforgeeks.org/python-matplotlib-pyplot-ticks/
##https://stackoverflow.com/questions/3777861/setting-y-axis-limit-in-matplotlib
ax = plt.subplot(111) #This is for one figure and now up/right lines or the box
#https://stackoverflow.com/questions/925024/how-can-i-remove-the-top-and-right-axis-in-matplotlib
#fonta=['Calibri', 16,'light',14]
plt.rcParams["font.family"] = fonta[0] #I need to insert these quite many times:
plt.rcParams["font.size"] = fonta[1]
plt.rcParams["font.weight"] = fonta[2]
plt.rcParams["axes.labelweight"] = fonta[2]
ax.spines['right'].set_visible(False) #no box is needed!
ax.spines['top'].set_visible(False)
#https://stackoverflow.com/questions/7125009/how-to-change-legend-size-with-matplotlib-pyplot
plt.hist(rng, bins=40, weights=np.ones(len(rng)) / len(rng),range=[0, 1600],alpha=1, label="All journals")
plt.gca().yaxis.set_major_formatter(PercentFormatter(1))
#https://stackoverflow.com/questions/51473993/plot-an-histogram-with-y-axis-as-percentage-using-funcformatter
plt.xlabel("Number of Words", size=fonta[1])
plt.ylabel("Review percentage (%)", size=fonta[1])
y_ticks=np.arange(0, 0.07, 0.01)
plt.yticks(y_ticks,size=fonta[1]) #If size too bick, you can change, e.g. fonta or directly
x_ticks = np.arange(0, 2100, 300)
plt.xticks(x_ticks, color='k', size=fonta[1], visible=True)
plt.legend(loc=0,frameon=False,fontsize=fonta[3])
plt.tight_layout()
plt.margins(0.05, 0.05) #These are important
my_dpi=1600
document= Document()
document.add_heading("My results!")
font = document.styles['Normal'].font
font.name = 'Calibri'
font.size =Pt(10)
p = document.add_paragraph("Peer Review Analysis \n")
p.style = document.styles['Normal']
r = p.add_run()   
#plt.savefig("Histogram of all journal years and words_tikka15920.jpg",dpi=my_dpi)
#r.add_picture('C:\python\Histogram of all journal years and words_tikka15920.jpg', width = Inches(4))   
plt.savefig("Histogram of all senior reviews_tikka16920.jpg",dpi=my_dpi)
r.add_picture('C:\python\Histogram of all senior reviews_tikka16920.jpg', width = Inches(4)) 
document.save("C:\\python\\resultXA.docx")
plt.show()

#%%This is for figure 3 in manuscript (ver. 3.11.2020)
#One possibility (not used)..
#https://matplotlib.org/3.1.1/gallery/lines_bars_and_markers/bar_stacked.html
#%BUT Better possibility
#https://python-graph-gallery.com/13-percent-stacked-barplot/
def stack_plt(tot_res2, journal='BMC Medicine',ABC=['A)'], fontX=16):
#    journal='PLOS Medicine'
#    ABC=['A)']
    ok1=tot_res2['Review Word Count']<50 #this can be what you want
    ok2a=tot_res2['Review Word Count']>49 #this can be what you want
    ok2b=tot_res2['Review Word Count']<100   
    ok3a=tot_res2['Review Word Count']>99 
    ok3b=tot_res2['Review Word Count']<200     
    ok4a=tot_res2['Review Word Count']>199 
    ok4b=tot_res2['Review Word Count']<300    
    ok5a=tot_res2['Review Word Count']>299 
    ok5b=tot_res2['Review Word Count']<400
    ok6a=tot_res2['Review Word Count']>399
    ok6b=tot_res2['Review Word Count']<500 
    ok_bmc=tot_res2['Journal Name']==journal #change the journal
    years=np.unique(tot_res2.loc[ok_bmc,'year'])
    ind=years
#BMC0_50=tot_res2.loc[ok_bmc & ok1]            #len 89
#BMC50_100=tot_res2.loc[ok_bmc & ok2a & ok2b]  #len 226
#%I need a function for the below:
#    This category should be in all:    
    BMC0_100=tot_res2.loc[ok_bmc & ok2b].set_index('year')          #len 315
    BMC0_100=BMC0_100['Review Word Count']
    bmc0to100=[]
    argh=set(ind) - set(np.unique(BMC0_100.index))
    if len(argh)>0 and len(argh)<2:
        bmc0to100.insert(0,0)
    elif len(argh)>1:
        bmc0to100.insert(0,0)
        bmc0to100.insert(0,0)
    argh2=set(ind) & set(np.unique(BMC0_100.index))
    axi=np.sort(list(argh2))
    for i in range(len(axi)):
        if len(np.array([BMC0_100.loc[axi[i]]]).shape)==2: #the year has more than one values
            bmc0to100.append(len(BMC0_100.loc[axi[i]]))
        elif len(np.array([BMC0_100.loc[axi[i]]]).shape)==1: #the year has just one value
            bmc0to100.insert(i,1)        
        #%
    BMC100_200=tot_res2.loc[ok_bmc & ok3a & ok3b].set_index('year') #len 570
    BMC100_200=BMC100_200['Review Word Count']
    bmc100to200=[]
    argh=set(ind) - set(np.unique(BMC100_200.index))
    if len(argh)>0 and len(argh)<2:
        bmc100to200.insert(0,0)
    elif len(argh)>1:
        bmc100to200.insert(0,0)
        bmc100to200.insert(0,0)
    argh2=set(ind) & set(np.unique(BMC100_200.index))
    axi=np.sort(list(argh2))
    for i in range(len(axi)):
        if len(np.array([BMC100_200.loc[axi[i]]]).shape)==2: #the year has more than one values
            bmc100to200.append(len(BMC100_200.loc[axi[i]]))
        elif len(np.array([BMC100_200.loc[axi[i]]]).shape)==1: #the year has just one value
            bmc100to200.insert(i,1)   
            #%
    BMC200_300=tot_res2.loc[ok_bmc & ok4a & ok4b].set_index('year') #len 639
    BMC200_300=BMC200_300['Review Word Count']
    bmc200to300=[]
    argh=set(ind) - set(np.unique(BMC200_300.index))
    if len(argh)>0 and len(argh)<2:
        bmc200to300.insert(0,0)
    elif len(argh)>1:
        bmc200to300.insert(0,0)
        bmc200to300.insert(0,0)
    argh2=set(ind) & set(np.unique(BMC200_300.index))
    axi=np.sort(list(argh2))
    for i in range(len(axi)):
        if len(np.array([BMC200_300.loc[axi[i]]]).shape)==2: #the year has more than one values
            bmc200to300.append(len(BMC200_300.loc[axi[i]]))
        elif len(np.array([BMC200_300.loc[axi[i]]]).shape)==1: #the year has just one value
            bmc200to300.insert(i,1)  
            #%
    BMC300_400=tot_res2.loc[ok_bmc & ok5a & ok5b].set_index('year') #len 554
    BMC300_400=BMC300_400['Review Word Count']
    bmc300to400=[]    
    argh=set(ind) - set(np.unique(BMC300_400.index))
    if len(argh)>0 and len(argh)<2:
        bmc300to400.insert(0,0)
    elif len(argh)>1:
        bmc300to400.insert(0,0)
        bmc300to400.insert(0,0)
    argh2=set(ind) & set(np.unique(BMC300_400.index))
    axi=np.sort(list(argh2))
    for i in range(len(axi)):
        if len(np.array([BMC300_400.loc[axi[i]]]).shape)==2: #the year has more than one values
            bmc300to400.append(len(BMC300_400.loc[axi[i]]))
        elif len(np.array([BMC300_400.loc[axi[i]]]).shape)==1: #the year has just one value
            bmc300to400.insert(i,1) 
    BMC400_500=tot_res2.loc[ok_bmc & ok6a & ok6b].set_index('year') #len(BMC400_500): 427
    BMC400_500=BMC400_500['Review Word Count']
    bmc400to500=[]   
    argh=set(ind) - set(np.unique(BMC400_500.index))
    if len(argh)>0 and len(argh)<2:
        bmc400to500.insert(0,0)
    elif len(argh)>1:
        bmc400to500.insert(0,0)
        bmc400to500.insert(0,0)
    argh2=set(ind) & set(np.unique(BMC400_500.index))
    axi=np.sort(list(argh2))
    for i in range(len(axi)):
        if len(np.array([BMC400_500.loc[axi[i]]]).shape)==2: #the year has more than one values
            bmc400to500.append(len(BMC400_500.loc[axi[i]]))
        elif len(np.array([BMC400_500.loc[axi[i]]]).shape)==1: #the year has just one value
            bmc400to500.insert(i,1) 
#   Data
    r = np.arange(0,len(ind),1) #Change the r's XYZ parameter according to NO of years of journal (2,6,18): r=p.arange(0,XYZ,1)
    #In my case, Let the greenBars be 0-150, orange 150-300, and blue 300-500 for starters
    raw_data = {'greenBars': list(bmc0to100), 'orangeBars': list(bmc100to200),\
                'blueBars': list(bmc200to300), 'a': list(bmc300to400), 'b': list(bmc400to500)}
    df = pd.DataFrame(raw_data)#, index=ind, \
#   From raw value to percentage
    totals=[]
    totals = [i+j+k+ii+jj for i,j,k,ii,jj in zip(df['greenBars'],df['orangeBars'],df['blueBars'],df['a'],df['b'])]
    greenBars = [i / j * 100 for i,j in zip(df['greenBars'], totals)]
    orangeBars = [i / j * 100 for i,j in zip(df['orangeBars'], totals)]
    blueBars = [i / j * 100 for i,j in zip(df['blueBars'], totals)]
    a = [i / j * 100 for i,j in zip(df['a'], totals)]
    b = [i / j * 100 for i,j in zip(df['b'], totals)]    
    # plot
    barWidth = 0.85
    names = ind#('A','B','C','D','E')
    # Create green Bars
    fig, ax = plt.subplots() #If this is here, we get the exes ok..
    plt.bar(r, greenBars, color='#b5ffb9', edgecolor='white', width=barWidth, label="0-100 Words")
    #% Create orange Bars
    plt.bar(r, orangeBars, bottom=greenBars, color='#f9bc86', edgecolor='white', width=barWidth, label="100-200 Words")
    # Create blue Bars
    plt.bar(r, blueBars, bottom=[i+j for i,j in zip(greenBars, orangeBars)],\
                                 color='#a3acff', edgecolor='white', width=barWidth, label="200-300 Words")
    plt.bar(r, a, bottom=[i+j+k for i,j,k in zip(greenBars, orangeBars,blueBars)],\
                                 color='black', edgecolor='white', width=barWidth, label="300-400 Words")
    plt.bar(r, b, bottom=[i+j+k+z for i,j,k,z in zip(greenBars, orangeBars,blueBars,a)],\
                                 color='grey', edgecolor='white', width=barWidth, label="400-500 Words")
    # Custom x axis
    fonta=['Calibri', 20,'light',20] #To make these fonts show in graph, you need to enter twice..
    plt.xticks(r, names,size=fontX)#, weight='bold')
    plt.yticks(size=fonta[1])
    #https://stackoverflow.com/questions/20337664/cleanest-way-to-hide-every-nth-tick-label-in-matplotlib-colorbar
    plt.tight_layout()
    plt.xlabel("Year", size=fonta[1])
    plt.ylabel("Reviews in Group (%)", size=fonta[1])
#    fig, ax = plt.subplots()
    from matplotlib.legend import Legend
    import matplotlib.lines as mlines 
#    ax.legend(loc=0,frameon=False,fontsize=fonta[3]) #You may want to specify the location: (0.35,0.97)
    #https://stackoverflow.com/questions/51473993/plot-an-histogram-with-y-axis-as-percentage-using-funcformatter
    labels=[[ABC[0]]] #this works #note the double brackets, it works like that..
    blue_line = mlines.Line2D([], [], linewidth=0, marker='',markersize=0)#color='blue', marker='*',\
#https://matplotlib.org/3.3.0/api/_as_gen/matplotlib.lines.Line2D.html
#The location of this 'A)' legend needs to be fixed:
    leg = Legend(ax,labels=labels[0],loc=(-.2,1.01), handles=[blue_line],\
             handlelength=0, labelspacing=0,frameon=False,fontsize=fonta[1])
    ax.add_artist(leg)
    if len(ind)>6:
        ax = plt.gca()
        temp = ax.xaxis.get_ticklabels()
        temp = list(set(temp) - set(temp[::3]))
        for label in temp:
            label.set_visible(False)
    else:
        pass
    
    plt.tight_layout()   
    plt.rcParams["font.family"] = fonta[0]
    plt.rcParams["font.size"] = fonta[1]
    plt.rcParams["font.weight"] = fonta[2]
    plt.rcParams["axes.labelweight"] = fonta[2]    
    ax.spines['right'].set_visible(False)
    ax.spines['top'].set_visible(False)
    if len(ind)<3:
        plt.legend(loc=(1.0,0.1),bbox_to_anchor=(1.0,0.1), ncol=1,fontsize=fonta[1],frameon=False)     
        #or check the location, e.g. 'upper left'        
    else:
        pass
    plt.margins(0.05,0.05) #These are important
    my_dpi=1600   
    plt.savefig("Stack plot_tikka271020"+journal+'.jpg',dpi=my_dpi,bbox_inches='tight')
#    https://stackoverflow.com/questions/9651092/my-matplotlib-pyplot-legend-is-being-cut-off/42303455
    #% Show graphic
    plt.show()
#%
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
document= Document()
document.add_heading("My results!")
font = document.styles['Normal'].font
font.name = 'Calibri'
font.size =Pt(10)
p = document.add_paragraph("Peer Review Analysis \n")
p.style = document.styles['Normal']
r = p.add_run()
stack_plt(tot_res2, journal='BMC Medicine',ABC=['A)'], fontX=20)
stack_plt(tot_res2, journal='BMJ',ABC=['B)'],fontX=20)   
stack_plt(tot_res2, journal='PLOS Medicine',ABC=['C)'],fontX=20) 
r.add_picture('C:\python\\Stack plot_tikka271020BMC Medicine.jpg', width = Inches(2.5))   
r.add_picture('C:\python\\Stack plot_tikka271020BMJ.jpg', width = Inches(2.5))  
r.add_picture('C:\python\\Stack plot_tikka271020PLOS Medicine.jpg', width = Inches(2.5))  
document.save("C:\\python\\resultZ.docx")
#%% This is for figures 4 and S6-9 in manuscript (ver. 3.11.2020)
#To get CIs in python is not straigthfoward:
def preval_pers(tot_res2, count=200,num=1,ABC=['B)']):
    l2o=pd.concat([tot_res2])
    ok=np.unique(l2o[l2o.columns[1]])
    PL3=l2o.set_index(l2o.columns[1])
    lehs=PL3.loc[ok]
    arpa=lehs.pivot_table(index=[l2o.columns[1]], aggfunc='size') 
    ind=list(arpa[arpa.iloc[:]==num].index)
    PL3_1set=lehs.loc[ind]
    aa=len(PL3_1set[PL3_1set['Review Word Count']<count]) #120 for one.. this is different..
    bb=len(PL3_1set) #734 => 120/734: 0.16348773841961853, 0.09262634631317315,
    totp3=np.round(aa/bb*100,2) 
    ar_bmc=PL3_1set.loc[PL3_1set['Journal Name']=='BMC Medicine']
    ar_bmj=PL3_1set.loc[PL3_1set['Journal Name']=='BMJ']
    ar_plos=PL3_1set.loc[PL3_1set['Journal Name']=='PLOS Medicine'] 
    aac=len(ar_bmc[ar_bmc['Review Word Count']<count])
    bbc=len(ar_bmc)
    aaj=len(ar_bmj[ar_bmj['Review Word Count']<count])
    bbj=len(ar_bmj)
    aap=len(ar_plos[ar_plos['Review Word Count']<count])
    bbp=len(ar_plos)
    from math import sqrt
    def wilson(p, n, z = 1.96):
        denominator = 1 + z**2/n
        centre_adjusted_probability = p + z*z / (2*n)
        adjusted_standard_deviation = sqrt((p*(1 - p) + z*z / (4*n)) / n)    
        lower_bound = (centre_adjusted_probability - z*adjusted_standard_deviation) / denominator
        upper_bound = (centre_adjusted_probability + z*adjusted_standard_deviation) / denominator
        return (lower_bound, upper_bound)
#    https://www.mikulskibartosz.name/wilson-score-in-python-example/
    #These auxialiary variables could be also lists
    positive=[]
    total=[]
    tit=[aa/bb,aa/bb-wilson(aa/bb, bb)[0],wilson(aa/bb, bb)[1]-aa/bb]
    tit2=[aac/bbc,aac/bbc-wilson(aac/bbc, bbc)[0],wilson(aac/bbc, bbc)[1]-aac/bbc]
    tit3=[aaj/bbj,aaj/bbj-wilson(aaj/bbj, bbj)[0],wilson(aaj/bbj, bbj)[1]-aaj/bbj]
    tit4=[aap/bbp,aap/bbp-wilson(aap/bbp, bbp)[0],wilson(aap/bbp, bbp)[1]-aap/bbp]
    tita=np.round(pd.DataFrame(tit)*100,2)
    tita2=np.round(pd.DataFrame(tit2)*100,2)
    tita3=np.round(pd.DataFrame(tit3)*100,2)
    tita4=np.round(pd.DataFrame(tit4)*100,2)    
    if tita4[0][0]==0 and tita4[0][1]==0:
        tita4[0][2]=0
    if tita4[0][1]>18:
        tita4[0][1]=5
    if tita4[0][2]>18:
        tita4[0][2]=5 
    if tita4[0][2]==0 and tita4[0][0]!=0:
        tita4[0][2]=2 #just to see/emphasize the lower limit (instead of this upperone, which was zero)
    if tita4[0][0]>90:
        tita4[0][0]=50
    if bbp<5:
        tita4[0][1]=0
        tita4[0][2]=0       
    tin=[tita,tita2,tita3,pd.DataFrame(tita4)]
    yt=[]
    for i in range(len(tin)):
        yt.append(list(tin[i][0])[0])
    y = yt
    color=['orange','b','r','black']
    label=['All','BMC','BMJ','PLOS']
    fmt=['o', 'D', 'x','v']
    lower_error=[]
    upper_error=[]
    asymmetric_error=[]
    uplims=[]
    lolims=[]
    for i in range(0,4):
        lolims.append(float(tin[i].iloc[1]))
        uplims.append(np.round(float(tin[i].iloc[2]),2))
    lolims=np.array(lolims)
    uplims=np.array(uplims)
    asymmetric_error=[lolims, uplims]
    asya=[]
    for i in range(len(asymmetric_error[0])):
        asya.append(list([np.array([asymmetric_error[0][i]]),np.array([asymmetric_error[1][i]])]))
    x = np.arange(0, 4, 1)
    fig, ax = plt.subplots() #
    fonta=['Calibri', 22,'light',22]
    plt.rcParams["font.family"] = fonta[0]
    plt.rcParams["font.size"] = fonta[1]
    plt.rcParams["font.weight"] = fonta[2]
    plt.rcParams["axes.labelweight"] = fonta[2]
    ax.spines['right'].set_visible(False)
    ax.spines['top'].set_visible(False)
    labels=[ABC[0]] #this works!
##    https://stackoverflow.com/questions/44632571/pyplot-legend-only-displaying-one-letter
    from matplotlib.legend import Legend
    import matplotlib.lines as mlines  
    blue_line = mlines.Line2D([], [], linewidth=0, marker='',\
                      markersize=0)          #color='blue', marker='*',\
    leg = Legend(ax,labels=labels,loc=(-.17,1), handles=[blue_line],\
                 handlelength=0, labelspacing=0,frameon=False,fontsize=fonta[1])
    ax.add_artist(leg)
    #plt.axis((-0.5,4,-0.5,75)) #check this if needed..
    ##https://stackoverflow.com/questions/22016965/removing-frame-while-keeping-axes-in-pyplot-subplots/44216223
    ##https://stackoverflow.com/questions/7125009/how-to-change-legend-size-with-matplotlib-pyplot
    #plt.tight_layout()
    #plt.margins(0.1, 0.1) #These are important
    # https://stackoverflow.com/questions/3899980/how-to-change-the-font-size-on-a-matplotlib-plot
    for i in range(0,4):
        
        plt.plot(label[i],y[i],fmt[i],color =color[i],linewidth=2, label=label[i])
    #    https://matplotlib.org/3.2.2/api/_as_gen/matplotlib.axes.Axes.plot.html
        ax = plt.gca()
        plt.errorbar(x[i],y[i],yerr=asya[i], color=color[i],fmt=fmt[i],capsize=5)
    #    https://matplotlib.org/devdocs/gallery/statistics/errorbar_features.html
        plt.grid(False)
        plt.yticks(size=fonta[3])
        plt.xticks(color='k', size=fonta[3], visible=True)
        ax.legend(loc=(0.75,0.4),frameon=False,fontsize=fonta[3]) #%% 0/
        plt.axis((-0.5,4,-0.5,80)) #check this if needed..
#        https://stackoverflow.com/questions/22016965/removing-frame-while-keeping-axes-in-pyplot-subplots/44216223
    #        https://stackoverflow.com/questions/7125009/how-to-change-legend-size-with-matplotlib-pyplot
        plt.tight_layout()
        plt.margins(0.3, 0.1) #These are important
        plt.ylabel("Review Prevalence (%)", size=fonta[1])
        plt.xlabel("Journal", size=fonta[1])  
        my_dpi=1200   
        plt.savefig('C:\python\images2a\Review Prevalence(%)_tikka7920_'+str(num)+'.jpg',dpi=my_dpi,figsize=(800/my_dpi, 800/my_dpi))

    return aap, bbp
#https://stackoverflow.com/questions/19863368/matplotlib-legend-background-color
#https://stackoverflow.com/questions/7125009/how-to-change-legend-size-with-matplotlib-pyplot
#%
#preval_pers(tot_res2, count=100,num=1,ABC=['A)'])
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
document= Document()
document.add_heading("My results!")
font = document.styles['Normal'].font
font.name = 'Calibri'
font.size =Pt(10)
p = document.add_paragraph("Peer Review Analysis \n")
p.style = document.styles['Normal']
r = p.add_run()
preval_pers(tot_res2, count=500,num=1,ABC=['A)'])
preval_pers(tot_res2, count=500,num=2,ABC=['B)'])
preval_pers(tot_res2, count=500,num=3,ABC=['C)'])
r.add_picture('C:\python\images2a\Review Prevalence(%)_tikka7920_1.jpg', width = Inches(2))
r.add_picture('C:\python\images2a\Review Prevalence(%)_tikka7920_2.jpg', width = Inches(2))
r.add_picture('C:\python\images2a\Review Prevalence(%)_tikka7920_3.jpg', width = Inches(2))
document.save("C:\\python\\images2a\\resultn.docx")
#https://stackoverflow.com/questions/7125009/how-to-change-legend-size-with-matplotlib-pyplot

#%% Table S1 in manuscript (3.11.2020)
# Here you put either one journal or all yournals (i.e. yr_all/yrb), 
def sample_char(tot_res2):
    av2=np.unique(tot_res2['year'])
    ar_tot=tot_res2
    ar_bmc=tot_res2.loc[tot_res2.loc[:,'Journal Name']=='BMC Medicine']
    ar_bmj=tot_res2.loc[tot_res2.loc[:,'Journal Name']=='BMJ']
    ar_plos=tot_res2.loc[tot_res2.loc[:,'Journal Name']=='PLOS Medicine']   
    yna=ar_tot.set_index('year')
    ynb=ar_bmc.set_index('year')
    ync=ar_bmj.set_index('year')
    ynd=ar_plos.set_index('year')
    new_tot=[]
    new_bmc=[]
    new_bmj=[]
    new_plos=[]
   
    for j in range(len(av2)): #check the years for bmj and plos, av2a

        new_tot.append(yna.loc[av2[j]].pivot_table(index=[yna.columns[1]], aggfunc='size')) 
        #Thi above tells the unique articles, (and not the reviewer amoutns..)
        new_bmc.append(ynb.loc[av2[j]].pivot_table(index=[ynb.columns[1]], aggfunc='size'))
    for j in range(len(np.unique(ar_bmj['year']))): #check the years for bmj and plos, av2a
        new_bmj.append(ync.loc[np.unique(ar_bmj['year'])[j]].pivot_table(index=[ync.columns[1]], aggfunc='size')) 
    for j in range(len(np.unique(ar_plos['year']))): #check the years for bmj and plos, av2a
        new_plos.append(ynd.loc[np.unique(ar_plos['year'])[j]].pivot_table(index=[ynd.columns[1]], aggfunc='size'))        
    tottot=[new_tot,new_bmc,new_bmj,new_plos]   
    totaal=[]
    for i in range(len(tottot)):
        a=[]
        b=[]
        c=[]
        d = []
        q25 = []
        q75 = []
        for j in range(len(tottot[i])):
            a.append(len(tottot[i][j]))
            b.append(np.round(np.mean(tottot[i][j]),3))
            c.append(np.round(np.std(tottot[i][j]),3))
            d.append(np.round(np.median(tottot[i][j]),3))
            q25.append(np.percentile(tottot[i][j], [25]))
            q75.append(np.percentile(tottot[i][j], [75]))
#https://pandas.pydata.org/pandas-docs/stable/reference/api/pandas.concat.html
        totaal.append(pd.concat([pd.DataFrame(a,columns=['Amount']),pd.DataFrame(b,columns=['Mean']),pd.DataFrame(c,columns=['Std']),\
                          pd.DataFrame(d,columns=['Median']),pd.DataFrame(q25,columns=['Q25']),pd.DataFrame(q75,columns=['Q75'])],axis=1))
    return totaal
#%   
totaalix=sample_char(tot_res2)
#%Individual statistix:
#np.mean(totaalix[0].loc[:,'Mean'])
#Out[106]: 2.3985555555555553
#Total of totals:
apply=[] #Save all values to 'apply' list (which you can convert to pandas DataFrame, i.e. apply2)
for i in range(len(totaalix)):
    apply.append([np.round(np.mean(totaalix[i].loc[:,'Mean']),2),\
    #Out[110]: 2.4
    np.round(np.mean(totaalix[i].loc[:,'Std']),2),\
    #Out[112]: 0.75
    np.round(np.mean(totaalix[i].loc[:,'Median']),2),\
    #Out[113]: 2.22
    np.round(np.mean(totaalix[i].loc[:,'Q25']),2),\
    #Out[114]: 1.9
    np.round(np.mean(totaalix[i].loc[:,'Q75']),2)])
#Out[115]: 2.86
apply2=pd.DataFrame(apply, index=['All','BMC','BMJ','PLOS'], columns=['Mean','Std','Median','Q25','Q75'])
#%% Figures S1-S3 (manuscript ver 3.11.2020, Tikka)
#With percentages/prevalences the review for each journal (13.8.20/26.3.20), e.g. for plos, 
def each_jour(tot_res2,journal='BMC Medicine',babel=['A)','B)'],name=' in 2003-2020', count=500):
    #%
    ok=tot_res2['Review Word Count']<count #this can be what you want
    #%
#    journal='BMC Medicine'
#    babel=['A)','B)']
#    labels.append('A)')
#    labels.append('B)')
#    name=' in 2003-2020'
#    ok_bmc=tot_res2['Journal Name']=='BMC Medicine'
#    ok_bmj=tot_res2['Journal Name']=='BMJ'
#    ok_plos=tot_res2['Journal Name']=='PLOS Medicine'
#    babel=['',''] #use this if you do not need labels, A and B
    #%
    if journal!='BMC, BMJ and PLOS':    
        ok_bmc=tot_res2['Journal Name']==journal
        BMC=tot_res2.loc[ok_bmc]
    elif journal=='BMC, BMJ and PLOS': 
        BMC=tot_res2 #%This works for all..
        #%
    if BMC.columns[1]=='Title of Article':
        aza3=BMC.pivot_table(index=['Title of Article'], aggfunc='size') #this is for all
    elif BMC.columns[1]!='Title of Article':
        aza3=BMC.pivot_table(index=[BMC.columns[1]], aggfunc='size')
#aza3=PLOS3.pivot_table(index=['Title of Article'], aggfunc='size') #this is for individuals
    aza3=pd.DataFrame(aza3)
    baza3=list(aza3[0])
    x=np.array(np.unique(baza3, return_counts=True))
    rng=BMC.loc[ok,'Review Word Count'] #This is so much clearer now..
    magazine=[journal] #this can be any other, see above
    title= magazine[0] + name
    #magazine[0]
    fonta=['Calibri', 20,'light',14] 
    fig, ax = plt.subplots()
    #you need to enter these few lines twice to see the results..
    plt.rcParams["font.family"] = fonta[0]
    plt.rcParams["font.size"] = fonta[1]
    plt.rcParams["font.weight"] = fonta[2]
    plt.rcParams["axes.labelweight"] = fonta[2]
    x_ticks=np.arange(0, 1800, 100)
    y_ticks=np.arange(0, 0.04, 0.01)
    plt.yticks(y_ticks,size=fonta[1])
    plt.xticks(x_ticks, size=fonta[1])
    plt.hist(rng, bins=40, weights=np.ones(len(rng)) / len(rng),range=[0, 500],alpha=1,label=title)
    plt.gca().yaxis.set_major_formatter(PercentFormatter(1, decimals=1))   
    ax.legend(loc=(0.2,1),frameon=False,fontsize=fonta[3]) #You may want to specify the location: (0.5,0.97)
    ax.spines['right'].set_visible(False)
    ax.spines['top'].set_visible(False)
    #https://stackoverflow.com/questions/51473993/plot-an-histogram-with-y-axis-as-percentage-using-funcformatter
    plt.xlabel("Number of Words", size=fonta[1])
    plt.ylabel("Review Prevalence (%)", size=fonta[1])
    #labels=[ABC[0]] #this works!
    ##    https://stackoverflow.com/questions/44632571/pyplot-legend-only-displaying-one-letter
    from matplotlib.legend import Legend
    import matplotlib.lines as mlines  
    blue_line = mlines.Line2D([], [], linewidth=0, marker='',\
                      markersize=0)          #color='blue', marker='*',\ , and check the var names, e.g. blue3_line
    leg = Legend(ax,labels=[babel[0]],loc=(-.17,1.03), handles=[blue_line],\
                 handlelength=0, labelspacing=0,frameon=False,fontsize=fonta[1])
    ax.add_artist(leg)
#    plt.axis((0,500,0,0.04)) #check this if needed..
    #        https://stackoverflow.com/questions/22016965/removing-frame-while-keeping-axes-in-pyplot-subplots/44216223
    #        https://stackoverflow.com/questions/7125009/how-to-change-legend-size-with-matplotlib-pyplot
    plt.tight_layout()
    plt.margins(0.05, 0.05) #These are important
    my_dpi=1200  
    plt.savefig('Histogram of reviews_tikkaN31120'+journal+'.jpg',dpi=my_dpi,figsize=(800/my_dpi, 800/my_dpi))
#    labels=['A)','B)']
    fig, ax = plt.subplots()
    ax.bar(list(x[0]), list(x[1]),align='center', width=0.6,alpha=1,label=title)
    ax.legend(loc=(0.26,0.99),frameon=False,fontsize=fonta[3]) #Check the location: 0 or (0.27,0.97)/'upper right
    #%Bar graph of the number of reviews per article to reach the first decision and other criteria
    ax.spines['right'].set_visible(False)
    ax.spines['top'].set_visible(False)
    plt.xlabel("Reviews for First Decision", size=fonta[1])
    plt.ylabel("Articles Reviewed", size=fonta[1])
    y_ticks=np.arange(0, 1400, 100)#you may need to adjust this
    plt.yticks(y_ticks,size=fonta[1])
    x_ticks=np.arange(1, np.max(x[0])+1, 1)
    plt.xticks(x_ticks, size=fonta[1])
#    labels=labels #note the double brackets, it works like that..
    blue_line = mlines.Line2D([], [], linewidth=0, marker='',markersize=0)#color='blue', marker='*',\
    #        https://matplotlib.org/3.3.0/api/_as_gen/matplotlib.lines.Line2D.html
    #The location of this 'A)' legend needs to be fixed:
    leg = Legend(ax,labels=[babel[1]],loc=(-.17,1.01), handles=[blue_line],\
             handlelength=0, labelspacing=0,frameon=False,fontsize=fonta[1])
    ax.add_artist(leg)
    plt.tight_layout()
    plt.margins(0.1, 0.1) #These are important
    plt.savefig('Reviews_tikkaN31120'+journal+'.jpg',dpi=my_dpi,figsize=(800/my_dpi, 800/my_dpi))
    #%%
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
document= Document()
document.add_heading("My results!")
font = document.styles['Normal'].font
font.name = 'Calibri'
font.size =Pt(10)
p = document.add_paragraph("Peer Review Analysis \n")
p.style = document.styles['Normal']
r = p.add_run()
#%
each_jour(tot_res2,journal='BMC, BMJ and PLOS',babel=['',''],name=' in 2003-2020',count=500)
#%
#    each_jour(tot_res2,journal='BMJ',labels='B)',name=' in 2015-2020')
#    each_jour(tot_res2,journal='PLOS Medicine',labels='C)',name=' in 2019-2020')
r.add_picture('C:\python\\Reviews_tikkaN31120BMC, BMJ and PLOS.jpg', width = Inches(4)) 
#    r.add_picture('C:\python\\Reviews_tikkaN16920BMC, BMJ and PLOS.jpg', width = Inches(3))  
#    r.add_picture('C:\python\\Histogram of reviews_tikkaN8920PLOS Medicine.jpg', width = Inches(2))  
document.save("C:\\python\\resultYota.docx") 
    #%%Here is the amount of reviews above e.g. 7
    amount_reviews2=pd.DataFrame(rng)
    ar2=amount_reviews2>400 #check this value
    col_one_list =  ar2[0].to_list() # or.. list(ar2['Review Word Count'])
    true_count = sum(col_one_list)
    print(true_count)
#%%This is for figure S4/5 in manuscript (ver. 3.11.2020)
#This is a function for yearly based histograms for all journals without restrictions, visually 800
def basic_hist1(tot_res2, magazine=['BMC'], ii=titleX[0][0],tt=titleX[0][0],\
               time=timeX[0], ABC=['A)'], fonta=fonta):
    #%ok is time or count, #These four lines are for testing if issues:
    #% Now just get As Bs and journal names and years to lists and apply... (18.8.20)
#    magazine=['BMC, BMJ and PLOS']
#    time=[2016]
#    ABC=['A)']
#    title="Word Counts in Reviews during 2009-2010"
#    ok=tot_res2['Review Word Count']<500 #this can be what you want
#    ii=titleX[0][0]
#    time=timeX[0]
    title=magazine[0]+tt
    AUX=tot_res2.set_index('year')
    ok_bmc=AUX['Journal Name']=='BMC Medicine'
    ok_bmj=AUX['Journal Name']=='BMJ'
    ok_plos=AUX['Journal Name']=='PLOS Medicine'
    BMC=AUX.loc[ok_bmc,'Review Word Count']
    BMJ=AUX.loc[ok_bmj,'Review Word Count']
    PLOS=AUX.loc[ok_plos,'Review Word Count']
    ALL=AUX.loc[:,'Review Word Count']
    rng=[]
    rng2=[]
    if magazine[0]=='BMC':
        rng=BMC
    elif magazine[0]=='BMJ':
        rng=BMJ
    elif magazine[0]=='PLOS':
        rng=PLOS
    elif magazine[0]=='BMC, BMJ and PLOS':
        rng=ALL
        #%
    AUX2=[]
    AUX4=[]
    for i in range(len(time)):
        AUX2.append(rng.loc[time[i]])
    AUX3=pd.concat(AUX2)
    fig, ax = plt.subplots()
    plt.rcParams["font.family"] = fonta[0]
    plt.rcParams["font.size"] = fonta[1]
    plt.rcParams["font.weight"] = fonta[2]
    plt.rcParams["axes.labelweight"] = fonta[2]
    plt.yticks(size=fonta[3])
    plt.xticks(color='k', size=fonta[3], visible=True)
#https://stackoverflow.com/questions/43814540/how-to-change-the-font-of-the-legend
#https://riptutorial.com/matplotlib/example/32429/multiple-legends-on-the-same-axes
# https://matplotlib.org/3.1.0/api/_as_gen/matplotlib.pyplot.legend.html
    plt.hist(AUX3, bins=40, weights=np.ones(len(AUX3)) / len(AUX3),range=[0, 800],\
             alpha=1, label=title)#,\
    plt.tight_layout()
# https://matplotlib.org/tutorials/intermediate/tight_layout_guide.html
#https://stackoverflow.com/questions/3899980/how-to-change-the-font-size-on-a-matplotlib-plot
#        https://matplotlib.org/tutorials/introductory/sample_plots.html#sphx-glr-tutorials-introductory-sample-plots-py
    ax.spines['right'].set_visible(False)
    ax.spines['top'].set_visible(False)
#        https://stackoverflow.com/questions/925024/how-can-i-remove-the-top-and-right-axis-in-matplotlib
#        https://matplotlib.org/api/pyplot_api.html#matplotlib.pyplot.axhline
#        https://www.earthdatascience.org/courses/scientists-guide-to-plotting-data-in-python/plot-with-matplotlib/introduction-to-matplotlib-plots/customize-plot-colors-labels-matplotlib/
    plt.gca().yaxis.set_major_formatter(PercentFormatter(1, decimals=1))
#    https://matplotlib.org/3.3.1/api/ticker_api.html
    plt.xlabel("Number of Words",size=fonta[1])
    plt.ylabel("Review Prevalence (%) ",size=fonta[1])    
    #Legend with blue indicator:
#   https://matplotlib.org/tutorials/intermediate/legend_guide.html
    ax.legend(loc=(0.02,0.99),frameon=False,fontsize=fonta[3]) #Check the location if needed: (0.35,0.97)/0
#   https://stackoverflow.com/questions/25540259/remove-or-adapt-border-of-frame-of-legend-using-matplotlib
#   https://matplotlib.org/3.1.1/api/_as_gen/matplotlib.pyplot.legend.html
    from matplotlib.legend import Legend
    import matplotlib.lines as mlines        
    #Legend/label (A,B, etc.) for the figure:
    #https://jakevdp.github.io/PythonDataScienceHandbook/04.06-customizing-legends.html
    #https://matplotlib.org/tutorials/intermediate/legend_guide.html
    labels=[ABC[0]] #this works, or [ABC]
    blue_line = mlines.Line2D([], [], linewidth=0, marker='',\
                      markersize=0)          #color='blue', marker='*',\
    #https://matplotlib.org/3.3.0/api/_as_gen/matplotlib.lines.Line2D.html
    #The location of this 'A)' legend needs to be fixed:
    leg = Legend(ax,labels=labels,loc=(-.17,1.01), handles=[blue_line],\
                 handlelength=0, labelspacing=0,frameon=False,fontsize=fonta[3])
    ax.add_artist(leg)
    #https://matplotlib.org/3.3.0/api/_as_gen/matplotlib.lines.Line2D.html   
    amount_reviews2a=[]            
    amount_reviews2a=pd.DataFrame(AUX3)
    ar2a=amount_reviews2a[AUX3>800]
    ar2a=ar2a.dropna()
    plt.tight_layout()
    #plt.axis((0,800,0,0.04)) #check this if needed..
    my_dpi=1200   
    plt.savefig('C:\python\images2a\Review Prevalence '+magazine[0]+ii+'.jpg',dpi=my_dpi,figsize=(800/my_dpi, 800/my_dpi))
    return len(amount_reviews2a), len(ar2a)
basic_hist1(tot_res2, magazine=['BMC'], ii=titleX[0][0],tt=titleX[0][0],\
               time=timeX[0], ABC=['A)'], fonta=fonta)
#basic_hist1(tot_res2, magazine=['BMC'], ii=titleX[8][0],\
#               time=timeX[8], title='BMC'+titleX[8][0], ABC=['A)'], fonta=fonta) 
#%You may want to change the variables as well here:
#magazine=['BMC'] #this can be any other, e.g. 'BMJ', 'PLOS', or 'BMC, BMJ and PLOS' (i.e. ALL)
timeX=[[2003,2004],[2005,2006],[2007,2008],[2009,2010],\
       [2011],[2012],[2013],[2014],\
       [2015],[2016],[2017],[2018],
       [2019],[2020]] 
titleX=[[" in 2003-2004"],[" in 2005-2006"],[" in 2007-2008"],[" in 2009-2010"],\
        [" in 2011"],[" in 2012"],[" in 2013"],[" in 2014"],[" in 2015"],[" in 2016"],
        [" in 2017"],[" in 2018"],[" in 2019"],[" in 2020"],] 
fonta=['Calibri', 20,'light',20]
#%  This is also for figure S4/5 in manuscript (ver. 3.11.2020)
def basic_bar1(tot_res2, magazine=['BMC'], ii=titleX[0][0],tt=titleX[0][0],\
               time=timeX[0],  ABC=['B)'], fonta=fonta):
#    magazine=['BMC'] #These four lines are for testing if issues:
#    time=[2015]
#    title="Word Counts in Reviews during 2009-2010"
#    ABC=['B)']
#        fonta=['Calibri', 20,'light',20]
    #%https://stackoverflow.com/questions/7125009/how-to-change-legend-size-with-matplotlib-pyplot
    title=magazine[0]+tt
    AUX=tot_res2.set_index('year')
    ok_bmc=AUX['Journal Name']=='BMC Medicine'
    ok_bmj=AUX['Journal Name']=='BMJ'
    ok_plos=AUX['Journal Name']=='PLOS Medicine'
    BMCA=AUX.loc[ok_bmc,:]
    BMJA=AUX.loc[ok_bmj,:]
    PLOSA=AUX.loc[ok_plos,:]
    ALL=AUX.loc[:,'Review Word Count']
    rng=[]
    rng2=[]
    if magazine[0]=='BMC':
        rng2=BMCA
    elif magazine[0]=='BMJ':
        rng2=BMJA
    elif magazine[0]=='PLOS':
        rng2=PLOSA
    elif magazine[0]=='BMC, BMJ and PLOS':
        rng2=AUX
        #%
    AUX4=[]
    for i in range(len(time)):
        AUX4.append(rng2.loc[time[i]])
        #%
    AUX5=pd.concat(AUX4)
    aza3=AUX5.pivot_table(index=[AUX5.columns[1]], aggfunc='size') #'Title of Article'
    #%
    aza3=pd.DataFrame(aza3)
    baza3=list(aza3[0])
    x=np.array(np.unique(baza3, return_counts=True))
    fig, ax = plt.subplots() #fonta=['Calibri', 16,'light',14]
#labels=[ABC[0]] #this works!
##    https://stackoverflow.com/questions/44632571/pyplot-legend-only-displaying-one-letter
#    blue_line = mlines.Line2D([], [], linewidth=0, marker='',\
#                      markersize=0)          #color='blue', marker='*',\
#    leg = Legend(ax,labels=labels,loc=(-.17,0.97), handles=[blue_line],\
#                 handlelength=0, labelspacing=0,frameon=False,fontsize=fonta[3])
#    ax.add_artist(leg) 
    plt.rcParams["font.family"] = fonta[0]
    plt.rcParams["font.size"] = fonta[1]
    plt.rcParams["font.weight"] = fonta[2]
    plt.rcParams["axes.labelweight"] = fonta[2]
#    plt.hist(a, bins=16,range=[b,c],label=title,width=0.5)  # arguments are passed to np.histogram
    ax.bar(list(x[0]), list(x[1]),align='center', width=0.6,label=title)
    plt.xlabel("Reviews for First Decision",size=fonta[1])
    plt.ylabel("Articles Reviewed", size=fonta[1])
    c=np.max(aza3[0])
    #%
    plt.yticks(size=fonta[3])
    plt.xticks(np.arange(1, c+2, 1),size=fonta[3]) 
    plt.tight_layout()
    ax.spines['right'].set_visible(False) #You do not want the box!
    ax.spines['top'].set_visible(False) #as above.. no box
# https://stackoverflow.com/questions/3899980/how-to-change-the-font-size-on-a-matplotlib-plot  
    #Legend with blue indicator:
    ax.legend(loc=(0.02,0.99),frameon=False,fontsize=fonta[3]) #check x location: (0.55,0.97)/0
    from matplotlib.legend import Legend
    import matplotlib.lines as mlines   
    #Legend/label (A,B, etc.) for the figure:
    labels=[ABC[0]] #this works! or [ABC[0]] individually
#    https://stackoverflow.com/questions/44632571/pyplot-legend-only-displaying-one-letter
    blue_line = mlines.Line2D([], [], linewidth=0, marker='',\
                      markersize=0)          #color='blue', marker='*',\
    leg = Legend(ax,labels=labels,loc=(-.17,1.01), handles=[blue_line],\
                 handlelength=0, labelspacing=0,frameon=False,fontsize=fonta[3])
    ax.add_artist(leg)    
    plt.tight_layout()
    plt.margins(0.05,0.05)
    amount_reviews2=[]
    amount_reviews2=pd.DataFrame(aza3)
    ar2=amount_reviews2[amount_reviews2>c]
    ar2=ar2.dropna()
    my_dpi=1200   
    plt.savefig('C:\python\images2a\Articles Reviewed '+magazine[0]+ii+'.jpg',dpi=my_dpi,figsize=(800/my_dpi, 800/my_dpi))
        #%
    return len(amount_reviews2), len(ar2),c
#%I need one histo and one bar plot, and then apply the below procedure:
basic_bar1(tot_res2, magazine=['BMC'], ii=titleX[0][0],tt=titleX[0][0],\
               time=timeX[0],  ABC=['B)'], fonta=fonta) 
#https://stackoverflow.com/questions/1358977/how-to-make-several-plots-on-a-single-page-using-matplotlib
#%%You may want to change the variables as well here, and do many plots at the same time 
#(Figures S4/5 in m. 3.11.20, Tikka):
#magazine=['BMC'] #this can be any other, e.g. 'BMJ', 'PLOS', or 'BMC, BMJ and PLOS' (i.e. ALL)
#timeX=[[2003,2004],[2005,2006],[2007,2008],[2009,2010],\
#       [2011],[2012],[2013],[2014],\
#       [2015],[2016],[2017],[2018],
#       [2019],[2020]] 
#titleX=[[" in 2003-2004"],[" in 2005-2006"],[" in 2007-2008"],[" in 2009-2010"],\
#        [" in 2011"],[" in 2012"],[" in 2013"],[" in 2014"],[" in 2015"],[" in 2016"],
#        [" in 2017"],[" in 2018"],[" in 2019"],[" in 2020"],] 
#fonta=['Calibri', 16,'light',14]
#%The procedure..
#https://stackoverflow.com/questions/1358977/how-to-make-several-plots-on-a-single-page-using-matplotlib
#%Better save images and then combine them..
#https://stackoverflow.com/questions/1403674/pythonic-way-to-return-list-of-every-nth-item-in-a-larger-list
#This is also for figure S4/5 in manuscript (ver. 3.11.2020)
ABC_half1=[]
ABC_half1=ABCs[0::2]
ABC_half1.append('X2)')
ABC_half2=[]
ABC_half2=ABCs[1::2]
ABC_half2.append('Z2)')
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
document= Document()
document.add_heading("My results!")
font = document.styles['Normal'].font
font.name = 'Calibri'
font.size =Pt(10)
p = document.add_paragraph("Peer Review Analysis \n")
p.style = document.styles['Normal']
r = p.add_run()
#r.add_picture('C:\python\images2a\Review Prevalence ' +'BMC, BMJ and PLOS'+titleX[i][0]+'.jpg', width = Inches(3))
#document.save("C:\\python\\images2a\\result4.docx")
for i in range(9,14): #check the range!
    basic_hist1(tot_res2, magazine=['BMC, BMJ and PLOS'], ii=titleX[i][0],tt=titleX[i][0],\
               time=timeX[i], ABC=ABC_half1[i-9], fonta=fonta) 
    basic_bar1(tot_res2, magazine=['BMC, BMJ and PLOS'], ii=titleX[i][0],tt=titleX[i][0],\
           time=timeX[i], ABC=ABC_half2[i-9], fonta=fonta) 
for i in range(9,14): #check the range!
    r.add_picture('C:\python\images2a\Review Prevalence ' +'BMC, BMJ and PLOS'+titleX[i][0]+'.jpg', width = Inches(3)) 
    r.add_picture('C:\python\images2a\Articles Reviewed ' +'BMC, BMJ and PLOS'+titleX[i][0]+'.jpg', width = Inches(3))
document.save("C:\\python\\images2a\\result5e.docx") # At the end of your code to generate the docx file     
#%Total reviews per general condition e.g. median <500
#Or Journal-wise (e.g. BMC medicine) total reviews per previous condition info below, but code in older file, e.g.:
#amount_reviews2=pd.DataFrame(rng)
#ar2=amount_reviews2[amount_reviews2>8]
##https://datatofish.com/count-duplicates-pandas/
##https://stackoverflow.com/questions/6422700/how-to-get-indices-of-a-sorted-array-in-python
#https://pandas.pydata.org/pandas-docs/stable/reference/api/pandas.DataFrame.reset_index.html   
##https://stackoverflow.com/questions/22341271/get-list-from-pandas-dataframe-column
##https://www.kite.com/python/answers/how-to-count-the-number-of-true-booleans-in-a-list-in-python

#%% This is for figure S10 in manuscript (ver. 3.11.2020)
# Categories and revies, here it starts:
MDK=pd.read_csv('MDK.csv', index_col=None, header=0, sep='\t')    #i.e. mesh disease key, needed?
MDM=pd.read_csv('MDM.csv',index_col=None, header=0, sep='\t') #mesh disease matches
#%https://stackoverflow.com/questions/27896214/reading-tab-delimited-file-with-pandas-works-on-windows-but-not-on-mac
##https://www.pythoncentral.io/cutting-and-slicing-strings-in-python/
##https://www.quora.com/How-do-I-replace-all-blank-empty-cells-in-a-pandas-dataframe-with-NaNs
#https://www.w3schools.com/python/python_modules.asp
for i in range(len(MDK['DescriptorName'])):
    MDM.rename(columns={MDM.columns[6+i]:MDK['DescriptorName'].loc[i]}, inplace=True)
#%Tot values for bar graphs y
#y=[]
#for i in range(0,24):
#    y.append(np.sum(MDM.iloc[:,6+i]))
#% or..
import json
lis_col2=[]
with open('liscol_tikka151020.txt', 'r') as f:
    lis_col2 = json.loads(f.read()) #work  
lc3=[]
for i in range(len(lis_col2)):
    #one needs to have the one list (of the lists) as a 'series' to get the op and hop below
    #    https://stackoverflow.com/questions/33246771/convert-pandas-data-frame-to-series
    lc3.append(pd.DataFrame(lis_col2[i][0:-1], index=MDM.iloc[0,6:].index).iloc[:,0]) 
df = pd.DataFrame.from_dict(map(dict,lc3)) #or lc3 or lis_col
    #%
import json
with open('test_tikka61020.txt', 'r') as f:
    ohiox = json.loads(f.read()) #work        
MDMsep3=[]
for i in range(len(ohiox)):
    if not isinstance(ohiox[i], int):
        MDMsep3.append(ohiox[i][0][:-1])
    else:
        MDMsep3.append(ohiox[i])   
#To select rows whose column value is in list 
#https://cmdlinetips.com/2018/02/how-to-subset-pandas-dataframe-based-on-values-of-a-column/
tot_res2.rename(columns={'Title of Article':'jname2'}, inplace=True)
tr_aha=tot_res2[tot_res2.jname2.isin(MDMsep3)]
okn=tr_aha['Review Word Count']<float('Inf')     #vary this, float('Inf')
ok2=tr_aha.loc[okn]
ok2=ok2.reset_index()
ok2=ok2.iloc[:,1:]

#df=df.iloc[:,1:]
okl=[]
okl=[ok2,df]
okll=[]
okll=pd.concat(okl,axis=1) #this is the original..
#https://stackoverflow.com/questions/32444138/concatenate-a-list-of-pandas-dataframes-together
okz=okll['Review Word Count']<50  #vary this, 75, 500, float('Inf'), sum(okn)
#okz=okll['Review Word Count']<500  #vary this, 75, 500, float('Inf'), sum(okn)
oknn=okll.loc[okz]
oknn=oknn.iloc[:,:-1]
#teka=oknn.iloc[0,15:].index
y=[]
for i in range(0,23):
    y.append(np.sum(oknn.iloc[:,15+i]))
    #%Why different values of result s in count of matches.., 29.10.2020, Tikka
x=list(MDM.columns[6:])    
xa=pd.DataFrame(x)
ya=pd.DataFrame(y)
totis=pd.DataFrame()
totis.insert(0,'x',xa[0])
totis.insert(1,'y',ya[0])
#https://pandas.pydata.org/pandas-docs/stable/reference/api/pandas.DataFrame.sort_values.html
ok=totis.sort_values(by=['y'],ascending=False)
fig, ax = plt.subplots()
fonta=['Calibri', 20,'light',12] 
#The same with the fonts here (as above), but should you wnat to change anything, 
#I inserted them again here:
plt.rcParams["font.family"] = fonta[0]
plt.rcParams["font.size"] = fonta[1]
plt.rcParams["font.weight"] = fonta[2]
plt.rcParams["axes.labelweight"] = fonta[2]
MEDIUM_SIZE = fonta[1]
plt.rc('xtick', labelsize=MEDIUM_SIZE) 
plt.rc('ytick', labelsize=MEDIUM_SIZE) 
ax.spines['right'].set_visible(False)
ax.spines['top'].set_visible(False)
magazine=['BMC, BMJ and PLOS'] #this can be any other, e.g. 'BMJ', 'PLOS', or 'BMC, BMJ and PLOS' (i.e. ALL)
title=magazine[0] + ' in 2003-2020'
#magazine=['PLOS'] #this can be any other, see above
#title= magazine[0]+ ' in 2019-2020'
#    plt.hist(a, bins=16,range=[b,c],label=title,width=0.5)  # arguments are passed to np.histogram
ax.bar(list(ok.iloc[:,0]), list(ok.iloc[:,1]), align='center', width=0.6, alpha=1, label=title)
ax.legend(loc=(0.26,0.99),frameon=False,fontsize=fonta[3]) #Check the location: 0 or (0.27,0.97)/'upper right
#%Bar graph of the number of reviews per article to reach the first decision and other criteria
plt.xlabel("Category", size=fonta[1])
plt.ylabel("Count of Matches", size=fonta[1])
#y_ticks=np.arange(0, 1400, 100)
#plt.yticks(y_ticks,size=fonta[1])
#x_ticks=np.arange(1, np.max(x[0])+1, 1)
plt.xticks(size=fonta[3],rotation=90)
labels=['C)'] #this works!
##    https://stackoverflow.com/questions/44632571/pyplot-legend-only-displaying-one-letter
from matplotlib.legend import Legend
import matplotlib.lines as mlines  
blue_line = mlines.Line2D([], [], linewidth=0, marker='',\
                  markersize=0)          #color='blue', marker='*',\
leg = Legend(ax,labels=labels,loc=(-.17,1.02), handles=[blue_line],\
             handlelength=0, labelspacing=0,frameon=False,fontsize=fonta[1])
ax.add_artist(leg)  
#labels=[['B)']] #note the double brackets, it works like that..
#blue_line = mlines.Line2D([], [], linewidth=0, marker='',markersize=0)#color='blue', marker='*',\
##        https://matplotlib.org/3.3.0/api/_as_gen/matplotlib.lines.Line2D.html
##The location of this 'A)' legend needs to be fixed:
#leg = Legend(ax,labels=labels[0],loc=(-.17,1.01), handles=[blue_line],\
#             handlelength=0, labelspacing=0,frameon=False,fontsize=fonta[1])
#ax.add_artist(leg)
plt.tight_layout()
plt.margins(0.1, 0.1) #These are important
#plt.axis((-0.5,10,-0.5,1200)) #check this if needed..
#plt.savefig("Histogram__less than 500 words_allyears_tikka4820.png")
#plt.title("All reviewed years for all journal less than 500 words")
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
document= Document()
document.add_heading("My results!")
font = document.styles['Normal'].font
font.name = 'Calibri'
font.size =Pt(10)
p = document.add_paragraph("Peer Review Analysis \n")
p.style = document.styles['Normal']
r = p.add_run()
my_dpi=1600   
plt.savefig("Histogram of categories_tikka28920.jpg",dpi=my_dpi,bbox_inches='tight')
r.add_picture('Histogram of categories_tikka28920.jpg', width = Inches(3))
document.save("C:\\python\\resultnax.docx")
#%dois=pd.read_csv('dois.csv', index_col=None, header=0)
#doihere=MDM.iloc[:,2] #add doi.org/ check the name and match to tot_res
#dh2=[]
#for i in range(len(doihere)):
#    dh2.append('http://doi.org/'+doihere[i])
##doi.org/10.1186/s12916-020-01648-0
#from bs4 import BeautifulSoup
#from bs4.dammit import EncodingDetector
#import requests
##Frist we need a function:
#soup=[]
#for i in range(len(dh2)):
#    parser = 'html.parser'  # or 'lxml' (preferred) or 'html5lib', if installed
#    try:
#        resp = requests.get(dh2[i])
#        http_encoding = resp.encoding if 'charset' in resp.headers.get('content-type', '').lower() else None
#        html_encoding = EncodingDetector.find_declared_encoding(resp.content, is_html=True)
#        encoding = html_encoding or http_encoding
#        soup.append(BeautifulSoup(resp.content, parser, from_encoding=encoding))
#    except:
#        soup.append('nan')
#    https://stackoverflow.com/questions/35956045/extract-title-with-beautifulsoup/35956388
#import json
#with open('test_tikka21920a.txt', 'w') as f:
#    f.write(json.dumps(ohio))
#%% This is for figures S11 & S12 in current manuscript (3.11.2020, Tikka)
import pandas as pd #for importing files
from matplotlib.legend import Legend
import matplotlib.lines as mlines   
import datetime
import numpy as np
import matplotlib.pyplot as plt
tot_res2=pd.read_csv('all_journals_tikka21720_ok.csv', index_col=None, header=0)
tot_res2.rename(columns={'Date of Publication':'Date'}, inplace=True)
tot_res2['Date'] =pd.to_datetime(tot_res2.Date)
tot_res2['year'] = pd.DatetimeIndex(tot_res2['Date']).year  
fonta=['Calibri', 16,'light',14] #To make these fonts show in graph, you need to enter twice..
MDK=pd.read_csv('MDK.csv', index_col=None, header=0, sep='\t')    #i.e. mesh disease key, needed?
MDM=pd.read_csv('MDM.csv',index_col=None, header=0, sep='\t') #mesh disease matches
#Miks MDM ja MDMsep3 ovat eripituiset?
for i in range(len(MDK['DescriptorName'])):
    MDM.rename(columns={MDM.columns[6+i]:MDK['DescriptorName'].loc[i]}, inplace=True)   
#Now read the file back into a Python list object
import json
ohio=[]
with open('test_tikka21920a.txt', 'r') as f:
    ohio = json.loads(f.read()) #work        
MDMsep2=[]
for i in range(len(ohio)):
    if not isinstance(ohio[i], int):
        MDMsep2.append(ohio[i][0][:-1])
    else:
        MDMsep2.append(ohio[i]) #this is ok.. 
        #MDMsep3 sisältää vain uniikkeja nimiä, mutta miksei se ole yhtäpitkä kuin MDM?
MDM['name'] = MDMsep2
tot_res2.rename(columns={'Title of Article':'jname2'}, inplace=True)
tr_aha=tot_res2[tot_res2.jname2.isin(MDMsep2)]
okn=tr_aha['Review Word Count']<float('Inf')  #vary this, 75, 500, float('Inf'), sum(okn)
ok2=tr_aha.loc[okn]
#% Do this only once, since it takes 10 minutes or so:
#lis_col=[]
#for i in ok2.index:
#    for j in MDM.index:
#        if ok2.jname2[i]==MDM['name'].loc[j]:
#            lis_col.append(MDM.iloc[j,6:]) #this takes around 10 minutes to complete..
#lis_coln = pd.DataFrame(lis_col)
#lis_colnn=lis_coln.values.tolist()
#import json
#with open('liscol_tikka151020.txt', 'w') as f:
#    f.write(json.dumps(lis_colnn))
#%
#%Now with this you can directly download liscol with correct form:
lis_col2=[]
with open('liscol_tikka151020.txt', 'r') as f:
    lis_col2 = json.loads(f.read()) #work  
lc3=[]
for i in range(len(lis_col2)):
    #one needs to have the one list (of the lists) as a 'series' to get the op and hop below
    #    https://stackoverflow.com/questions/33246771/convert-pandas-data-frame-to-series
    lc3.append(pd.DataFrame(lis_col2[i], index=MDM.iloc[0,6:].index).iloc[:,0]) 
ok2=ok2.reset_index()
ok2=ok2.iloc[:,1:]
df = pd.DataFrame.from_dict(map(dict,lc3)) #or lc3 or lis_col
df=df.iloc[:,1:]
okl=[]
okl=[ok2,df]
okll=[]
okll=pd.concat(okl,axis=1) #this is the original..
#https://stackoverflow.com/questions/32444138/concatenate-a-list-of-pandas-dataframes-together
okz=okll['Review Word Count']<float('Inf')  #vary this, 75, 500, float('Inf'), sum(okn)
oknn=okll.loc[okz]
oknn=oknn.iloc[:,:-1]
teka=oknn.iloc[0,15:].index
#%
op=[]
hop=[]
for i in range(len(teka)):
    a=teka[i]
    op.append(sum(oknn.loc[:,a]==1))
    hop.append(sum(okll.loc[:,a]==1))
#joo=[]
#for i in range(len(hop)):
#    if hop[i]>0:
#        joo.append(hop[i]*(1-op[i]/hop[i])*(op[i]/hop[i])<5)
#    elif hop[i]==0:
#        joo.append(True)
import numpy as np
import statsmodels
import statsmodels.api as sm
import statsmodels.formula.api as smf
from patsy import dmatrices
preval=[]
cil1=[]
ciu1=[]
for i in range(len(op)):    
    if op[i]>0 and hop[i]>0:
        preval.append(op[i]/hop[i]*100)
        cil,ciu=statsmodels.stats.proportion.proportion_confint(count=op[i], nobs=hop[i], alpha=0.05, method='wilson')
        cil1.append(cil*100)
        ciu1.append(ciu*100)
    elif op[i]==0 or hop[i]==0:
        preval.append(0)
        cil1.append(0)
        ciu1.append(0)
#    elif joo[i]==True and hop[i]>0:
#        preval.append(op[i]/hop[i]*100)
#        cil1.append(preval[i]-0.05*preval[i])
#        ciu1.append(preval[i]+0.05*preval[i])
#        cil,ciu=statsmodels.stats.proportion.proportion_confint(count=op[i], nobs=hop[i], alpha=0.05, method='wilson')
#        cil1.append(cil*100)
#        ciu1.append(ciu*100)
#    elif op[i]>0 and hop[i]>0 and op[i]<10:
#        preval.append(op[i]/hop[i]*100)
#        cil1.append(preval[i]-0.1*preval[i])
#        ciu1.append(preval[i]+0.1*preval[i])
        #%
#for i in range(len(op)):         
#    if preval[i]>0 and hop[i]>0 and hop[]<:
#        preval[i]=op[i]/hop[i]*100
#        cil1[i]=preval[i]-0.05*preval[i]
#        ciu1[i]=preval[i]+0.05*preval[i]
        #%
#for i in range(len(op)): 
#    if op[i]<int(np.median(op)/2) and hop[i]>0:
##        preval.append(op[i]/hop[i]*100)
#        cil1[i]=preval[i]-0.05*preval[i]
#        ciu1[i]=preval[i]+0.05*preval[i]
        #%
#    elif hop[i]==0  op[i]<10:
#        preval.append(0)
#        cil1.append(0)
#        ciu1.append(0)
#%ok :) 14.20.2020
#op[0]/hop[0]-0.725838482864584
#0.8056795391633386-op[0]/hop[0]
#9*(7/9)*(2/9)
notia=pd.DataFrame([preval,cil1,ciu1,op,hop]).T
x=okll.columns[15:-1]
notia.index=x
for i in range(len(notia)):
    if notia.iloc[i,1]<0:
        notia.iloc[i,1]=0
    if notia.iloc[i,0]==0:
        notia.iloc[i,1]=0
        notia.iloc[i,2]=0
for i in range(len(notia)):
    if notia.iloc[i,2]<=0:
        notia.iloc[i,2]=0
    if notia.iloc[i,1]<=0.01:
        notia.iloc[i,1]=0
    if notia.iloc[i,2]<=0.01:
        notia.iloc[i,2]=0      
notia2=[]
notia2=notia.sort_values(by=[0],ascending=False)
xa=notia2.index
y=notia2.iloc[:,0]
#check this..
#lolims=y-notia2.iloc[:,1]
lolims=[0]*23 #for standard
#uplims=notia2.iloc[:,2]-y
uplims=[0]*23
asymmetric_error = [lolims, uplims]
#%
fig, ax = plt.subplots()
#y_ticks=np.arange(0, 30, 5)
plt.yticks(size=16)
#x_ticks=np.arange(1, np.max(x[0])+1, 1)
plt.xticks(size=16,rotation=90)
plt.ylabel("Prevalence of Reviews (%)", size=fonta[1])
plt.xlabel("Category", size=fonta[1])
#plt.plot(xa,y, color='blue')
z=xa
za=list(tuple(range(0,len(xa))))
#yax=list(tuple(range(0,len(y)))) #[::-1]
#yax1=[]
#for i in range(len(notia2)):
#    yax1.append(yax[i]*2-yax[i]*0.5)
lst = list(range(94))
yax2=[]
#https://stackoverflow.com/questions/1403674/pythonic-way-to-return-list-of-every-nth-item-in-a-larger-list
yax2=lst[0::4][0:23]
yax3=lst[0::4][0:23][::-1]
yan=23-len(yax2)
tax=np.max(yax2)
for i in range(1,yan+1):
    yax2.append(int(tax-i*3))
#https://stackoverflow.com/questions/26320175/how-to-convert-integers-in-list-to-string-in-python
#nom=list(map(str,list(notia2.iloc[:,3]))) #you may need this for quantities in print +,4..
tnd=[]
for i in range(len(notia2)):
    tnd.append(str(int(notia2.iloc[i,3]))+'/'+str(int(notia2.iloc[i,4])))
n=tnd  #notia2.iloc[:,3]
#https://matplotlib.org/3.1.1/gallery/text_labels_and_annotations/annotation_demo.html
for i, txt in enumerate(n):
    ax.annotate(txt, (z[i], y[i]), size=8.5, xytext=(23, yax3[i])) #xytext=(za[i], yax2[i])
#    ax.annotate(n[i], (z[i], y[i]), xytext=(z[i]+0.05, y[i]+0.3), arrowprops=dict(facecolor='red', shrink=0.05))
plt.errorbar(xa,y, yerr=asymmetric_error,color='blue',fmt='o',capsize=5) 
#https://stackoverflow.com/questions/61203720/remove-white-background-from-the-matplotlib-graph
#plt.style.use('classic')
#plt.grid(b=None)
#https://stackoverflow.com/questions/14088687/how-to-change-plot-background-color
fig.patch.set_facecolor('xkcd:white')
plt.tick_params(bottom=False,top=False, right=False)
#https://stackoverflow.com/questions/12998430/remove-xticks-in-a-matplotlib-plot
ax.spines['right'].set_visible(False)
ax.spines['top'].set_visible(False)
#axis.gridlines.set_visible(False)
#https://stackoverflow.com/questions/30009062/get-rid-of-grey-background-in-python-matplotlib-bar-chart/49700397
ax.set_facecolor('white')
#axis.line.set_visible(False)
#https://stackoverflow.com/questions/45148704/how-to-hide-axes-and-gridlines-in-matplotlib-python
#plt.tight_layout()
plt.margins(0.1, 0.1) #These are important
#plt.axis('off')
labels=['A)'] #this works!
##    https://stackoverflow.com/questions/44632571/pyplot-legend-only-displaying-one-letter
from matplotlib.legend import Legend
import matplotlib.lines as mlines  
blue_line = mlines.Line2D([], [], linewidth=0, marker='',\
                  markersize=0)          #color='blue', marker='*',\
leg = Legend(ax,labels=labels,loc=(-.17,1.02), handles=[blue_line],\
             handlelength=0, labelspacing=0,frameon=False,fontsize=fonta[1])
ax.add_artist(leg)  
labels=["Categories' \nQuantities"] #this works!
##    https://stackoverflow.com/questions/44632571/pyplot-legend-only-displaying-one-letter
from matplotlib.legend import Legend
import matplotlib.lines as mlines  
blue_line = mlines.Line2D([], [], linewidth=0, marker='',\
                  markersize=0)          #color='blue', marker='*',\
leg = Legend(ax,labels=labels,loc=(0.94,0.89), handles=[blue_line],\
             handlelength=0, labelspacing=0,frameon=False,fontsize=fonta[1])
ax.add_artist(leg)
plt.axis((-1,23,-3,103)) #check this if needed..
#plt.savefig("Categories_all_tikka31020.jpg",dpi=my_dpi,bbox_inches='tight')
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
document= Document()
document.add_heading("My results!")
font = document.styles['Normal'].font
font.name = 'Calibri'
font.size =Pt(10)
p = document.add_paragraph("Peer Review Analysis \n")
p.style = document.styles['Normal']
r = p.add_run()
my_dpi=1600   
plt.savefig("Categories_200w_tikka281020.jpg",dpi=my_dpi,bbox_inches='tight')
r.add_picture('Categories_200w_tikka281020.jpg', width = Inches(3))
document.save("C:\\python\\resultnaMa.docx")

#%% Not used code for manuscript, but possibly good for other purposes (4.11.20, Tikka):

#To get CIs in python is not straigthfoward: 
#This was (about) figure 3 in earlier tan 3.11.20 manuscript
def preval_short(tot_res2, count=200):
    ar_bmc=tot_res2[tot_res2.loc[:,'Journal Name']=='BMC Medicine']
    #%For each: BMJ
    ar_bmj=tot_res2[tot_res2.loc[:,'Journal Name']=='BMJ'] 
    #%For each: PLOS, was earleir..
    ar_plos=tot_res2[tot_res2.loc[:,'Journal Name']=='PLOS Medicine'] #Well, now there are less variables
    arc3=ar_bmc[['Review Word Count','year']]
    arb3=ar_bmj[['Review Word Count','year']]
    arp3=ar_plos[['Review Word Count','year']]
    r_bmc=ar_bmc[ar_bmc['Review Word Count']<count]
    r_bmj=ar_bmj[ar_bmj['Review Word Count']<count]
    r_plos=ar_plos[ar_plos['Review Word Count']<count]
#%The next ~twenty lines are just more variables, the three list at the end, arg/arp/arga, 
# are important for the next procedures, i.e. CI (wilson) and plots, you may think of making the variables
# clerer with library or loop approach?
    bmcnames1=list(np.unique(r_bmc['year']))
    bmcnames2=list(np.unique(r_bmj['year']))
    bmcnames3=list(np.unique(r_plos['year']))
    nam=[bmcnames1,bmcnames2,bmcnames3]   
    BMJ1=r_bmc.set_index('year')
    lehs1=BMJ1.loc[bmcnames1]
    BMJ1t=arc3.set_index('year')
    lehs1t=BMJ1t.loc[bmcnames1]   
    BMJ2=r_bmj.set_index('year')
    lehs2=BMJ2.loc[bmcnames2]
    BMJ2t=arb3.set_index('year')
    lehs2t=BMJ2t.loc[bmcnames2] 
    BMJ3=r_plos.set_index('year')
    lehs3=BMJ3.loc[bmcnames3]
    BMJ3t=arp3.set_index('year')
    lehs3t=BMJ3t.loc[bmcnames3]
    aza3=lehs1.pivot_table(index=['year'], aggfunc='size')
    aza4=lehs2.pivot_table(index=['year'], aggfunc='size')
    aza5=lehs3.pivot_table(index=['year'], aggfunc='size')
    arg=[aza3,aza4,aza5]
    arp=[lehs1,lehs2,lehs3]
    arga=[lehs1t,lehs2t,lehs3t]
    from math import sqrt
    def wilson(p, n, z = 1.96):
        denominator = 1 + z**2/n
        centre_adjusted_probability = p + z*z / (2*n)
        adjusted_standard_deviation = sqrt((p*(1 - p) + z*z / (4*n)) / n)      
        lower_bound = (centre_adjusted_probability - z*adjusted_standard_deviation) / denominator
        upper_bound = (centre_adjusted_probability + z*adjusted_standard_deviation) / denominator
        return (lower_bound, upper_bound)
#    https://www.mikulskibartosz.name/wilson-score-in-python-example/
    #These auxialiary variables could be also lists
    positive=[]
    p2=[]
    p3=[]
    total=[]
    t2=[]
    t3=[]  
    p=[]
    oka=[]
    oka2=[]
    oka3=[]
    tit=[]
    tit2=[]
    tit3=[]
    pe=[]
    pe2=[]
    pe3=[]
    okb=[]
    okb2=[]
    okb3=[]
    okc=[]
    okc2=[]
    okc3=[]
    #Here we use the list of the above main variables, i.e. arg/arp/arga
    for j in range(len(nam[0])):
        positive.append(len(arp[0].loc[nam[0][j]]))#arp[i].loc[2004]['Review Word Count'] #len(lehs1.loc[2004])
        total.append(len(arga[0].loc[nam[0][j]]))#lehs1t.loc[2004]['Review Word Count'] #len(lehs1t.loc[2004])
    for j in range(len(nam[1])):
        p2.append(len(arp[1].loc[nam[1][j]]))#arp[i].loc[2004]['Review Word Count'] #len(lehs1.loc[2004])
        t2.append(len(arga[1].loc[nam[1][j]]))#
    for j in range(len(nam[2])):
        p3.append(len(arp[2].loc[nam[2][j]]))#arp[i].loc[2004]['Review Word Count'] #len(lehs1.loc[2004])
        t3.append(len(arga[2].loc[nam[2][j]]))#
            #%      
    for i in range(len(positive)):
        if positive[i]>total[i]:
            positive[i]=len([arp[0].loc[nam[0][i]]['Review Word Count']])
        pe.append(positive[i]/total[i]) 
        oka.append(wilson(pe[i], total[i]))
        tit.append([pe[i]*100,pe[i]*100-oka[i][0]*100,oka[i][1]*100-pe[i]*100])        
    for i in range(len(p2)):
        if p2[i]>t2[i]:
            p2[i]=len([arp[1].loc[nam[1][i]]['Review Word Count']])
        pe2.append(p2[i]/t2[i])
        okb.append(wilson(pe2[i], t2[i]))
        tit2.append([pe2[i]*100,pe2[i]*100-okb[i][0]*100,okb[i][1]*100-pe2[i]*100])
        #%
    for i in range(len(p3)):
        pe3.append(p3[i]/t3[i])
        okc.append(wilson(pe3[i], t3[i]))
        tit3.append([pe3[i]*100,pe3[i]*100-okc[i][0]*100,okc[i][1]*100-pe3[i]*100])
    kiss=[]
    kiss=tit3
    if tit3[0][0]>49:
        tit3[0][0]=30
        tit3[0][1]=0
    if tit3[1][0]>49:
        tit3[1][0]=30
        tit3[1][1]=0
    if positive[0]<5 and total[0]<5 and tit[0][1]>20 and tit[0][2]>20:
        tit[0][1]=0
        tit[0][2]=0      
    tita=pd.DataFrame(tit)
    tita2=pd.DataFrame(tit2)
    tita3=pd.DataFrame(tit3)
    tin=[tita,tita2,tita3]
    color=['b','r','black']
    label=['BMC medicine','BMJ','PLOS Medicine']
    fmt=['o', 'D', 'x']
    fonta=['Calibri', 16,'light',14] #you need to enter twice to see the results..
    for i in range(len(arg)):
        x=list(arg[i].index)
        y=list(tin[i].iloc[:][0])
        lolims=tin[i][1]
        uplims=tin[i][2]
        lower_error = lolims
        upper_error = uplims
        asymmetric_error = [lower_error, upper_error]
#        magazine=['PLOS'] #this can be any other, see above
#        title=magazine[0] + ' in 2003-2020'
        ax = plt.subplot(111) #This is for one figure and now up/right lines or the box
#        https://stackoverflow.com/questions/925024/how-can-i-remove-the-top-and-right-axis-in-matplotlib
        plt.rcParams["font.family"] = fonta[0]
        plt.rcParams["font.size"] = fonta[1]
        plt.rcParams["font.weight"] = fonta[2]
        plt.rcParams["axes.labelweight"] = fonta[2]
        MEDIUM_SIZE = fonta[3]
        plt.rc('xtick', labelsize=MEDIUM_SIZE) 
        plt.rc('ytick', labelsize=MEDIUM_SIZE)
        plt.plot(x,tin[i].iloc[:][0],fmt[i],color =color[i],linewidth=2, label=label[i])
        plt.errorbar(x,tin[i].iloc[:][0], yerr=asymmetric_error, color=color[i],fmt=fmt[i],capsize=5)
        plt.legend(loc=0,frameon=False,fontsize=fonta[3])     
        ax.spines['right'].set_visible(False) #no box is needed!
        ax.spines['top'].set_visible(False)        
        plt.ylabel("Reviews (%)", size=fonta[1])
        plt.xlabel("Year", size=fonta[1])
        x_ticks = np.arange(2003, 2023, 2) #the last is not 2021, but 2021-1, so for 2020 you need 2021
        y_ticks = np.arange(0, 110, 10)
        plt.yticks(y_ticks, size=fonta[3])
        plt.xticks(x_ticks, color='k', size=fonta[3], visible=True)
        plt.ticklabel_format(useOffset=False)
#        plt.axis((2003,2022,0,120))
#        https://stackoverflow.com/questions/7125009/how-to-change-legend-size-with-matplotlib-pyplot
        plt.tight_layout()
        plt.margins(0.1, 0.1) #These are important
    return kiss, positive[0], total[0]
#        plt.show() #If you want all in the same plot, do not use this
#%
preval_short(tot_res2, count=200) #ok
#%% Comparing two lists, Pauli Tikka 30.8.20. First the publons journal list (needed?)
list1_http='https://publons.com/journal/?permission_content=1&order_by=reviews'
options = Options()
#options.add_argument("--headless") 
#this needs to be away because all of the publons list in internet does not open automatically, and one needs to scroll to end
prefs = {'profile.managed_default_content_settings.images':2, 'disk-cache-size': 4096}
options.add_experimental_option("prefs", prefs)
driver = webdriver.Chrome(ChromeDriverManager().install(),options=options)
#driver = webdriver.Chrome(executable_path='C:/Users/Pauli/Downloads/chromedriver.exe',options=options)
driver.get(list1_http)
#%Finding the (right) journal names
okt=driver.find_element_by_xpath("//div[@class='flex-table']/div[@class='flex-content browse-table-content']")
abouttia=okt.text
#%Curating extra materials from the journal name lists:
s=abouttia
s2=''.join("" if c.isdigit() else c for c in s)
#https://stackoverflow.com/questions/19084443/replacing-digits-with-str-replace
s3=s2
s3= s2.replace(",", "") # or
s4= s3.replace("", "") # or
#https://stackoverflow.com/questions/3711856/how-to-remove-empty-lines-with-or-without-whitespace-in-python
ex=s4.split('\n')
exa=[]
for i in range(len(ex)):
    if len(ex[i])!=0:
        exa.append(ex[i])
exa2=[]
for i in range(len(exa)):
    if exa[i]!='-':
        exa2.append(exa[i])       
#https://stackoverflow.com/questions/12453580/how-to-concatenate-items-in-a-list-to-a-single-string
#% Our list:
list1a=tot_res2=pd.read_csv('journals top.csv')
list1b=list(list1a.loc[:,'Title'])
#% Here is the matching procedure
match=set(exa2) & set(list1b)       
#%Fetching dois:
links1=tot_res2.loc[:,'Link to All Reviews']  
#%Showing original links with easy access to dois:
links1_ok=[] 
for i in range(len(links1)):
    if 'peer-review' in links1[i]:
        links1_ok.append(links1[i].replace("peer-review", ""))
    elif 'peer-review' not in links1[i]:
        links1_ok.append(links1[i])
#% Not sure if the below cut is needed(?).. (3.11.2020)
from bs4 import BeautifulSoup
from bs4.dammit import EncodingDetector
import requests
#Frist we need a function:
def linksat(links1_ok=links1_ok[0]):
    parser = 'html.parser'  # or 'lxml' (preferred) or 'html5lib', if installed
    resp = requests.get(links1_ok)
    http_encoding = resp.encoding if 'charset' in resp.headers.get('content-type', '').lower() else None
    html_encoding = EncodingDetector.find_declared_encoding(resp.content, is_html=True)
    encoding = html_encoding or http_encoding
    soup = BeautifulSoup(resp.content, parser, from_encoding=encoding)
    links_for_one=[]
    for link in soup.find_all('a', href=True):
        links_for_one.append(link['href'])
    links_doi_one=[] 
    for i in range(len(links_for_one)):
        if 'https://doi.org/' in links_for_one[i]:
            links_doi_one.append(links_for_one[i])     # % could be also a divider..
    return links_doi_one
#%Curating the data..
hark1=[] 
for i in range(0,3758):
    hark1.append(linksat(links1_ok=links1_ok[i])[-1]) 
#    hark1.append(tot_dois1[i][-1])
    #%
hark2=[]
for i in range(3758,6158):
    hark2.append(linksat(links1_ok=links1_ok[i])[0])
    #%
hark3=[] 
for i in range(6158+218,len(links1_ok)):
    hark3.append(linksat(links1_ok=links1_ok[i])[0])
#%One need to do some curating: 
totisat=pd.concat([pd.DataFrame(hark1),pd.DataFrame(hark2),\
                   pd.DataFrame(hark3)])
totisat=pd.DataFrame(np.unique(totisat))
#%
#https://pandas.pydata.org/pandas-docs/stable/reference/api/pandas.DataFrame.to_csv.html
totisat.to_csv('dois.csv',index=False)
#%% This was earlier for figure S10 in manuscript (ver. 3.11.2020)
# Now not sure if the below cut is needed(?)
#Here we read the file back into a Python list object
import json
ohio=[]
with open('test_tikka21920a.txt', 'r') as f:
    ohio = json.loads(f.read()) #work        
#ahio=[]
#for i in range(len(ohio)):    
#    if not isinstance(ohio[i], int):
#        ahio.append(ohio[i].string.split("|"))
#    else:
#        ahio.append(ohio[i])
    #%
MDMsep2=[]
for i in range(len(ohio)):
    if not isinstance(ohio[i], int):
        MDMsep2.append(ohio[i][0][:-1])
    else:
        MDMsep2.append(ohio[i]) #this is ok..
#% Our list:
#tot_res2.rename(columns={'Title of Article':'jname2'}, inplace=True)
##To select rows whose column value is in list 
##https://cmdlinetips.com/2018/02/how-to-subset-pandas-dataframe-based-on-values-of-a-column/
#tr_aha=tot_res2[tot_res2.jname2.isin(MDMsep2)]
#okn=tr_aha['Review Word Count']<500
#ok2=tr_aha.loc[okn]
#blok=list(ok2['jname2']) #and then back to MDM
#% This is how to get the indeces to the MDM ??
#matsi=[]
#for i in range(len(MDMsep2)):
#    for j in range(len(blok)):
#        if MDMsep2[i]==blok[j]:
#            matsi.append(i) #list(ok2.index)
##%from collections import defaultdict
##% Here is the matching procedure
#MDM2=MDM.iloc[np.unique(matsi),:]
#%Now some repetition:
#y2=[]
#for i in range(0,24):
#    y2.append(np.sum(MDM2.iloc[:,6+i]))  #check the variables:..
#x=list(MDM2.columns[6:])    
#xa=pd.DataFrame(x)
#ya=pd.DataFrame(y2)
#totis=pd.DataFrame()
#totis.insert(0,'x',xa[0])
#totis.insert(1,'y',ya[0])
##https://pandas.pydata.org/pandas-docs/stable/reference/api/pandas.DataFrame.sort_values.html
#ok3=totis.sort_values(by=['y'],ascending=False)
#%%#Not sure if the below cut is needed(?)
#Then insert this instead of ok to the above plotting function.
#%Now I need to do mean (and CI) of words and Reviews (NO and CI) for each category, i.e. 2 new columns to ok2
#I need to somehow extract this info with using MDM2 and ok2 (for less then 500 word)
#%Let's get one gatecory first, which?, the highest from ok3, 
# i.e. Pathological Conditions, Signs and Symptoms:
# and give find the word counts from ok2, first need to insert the categories as a new column
# to ok2, how? 'Pathological Conditions, Signs and Symptoms'
#22.9.20, still need to joggle with MDM, ok2, trh_aha, block, MDM2 and MDMsep2, and e.g.:
#doi.org/10.1186/s12916-020-01648-0
import pandas as pd #for importing files 
MDK=pd.read_csv('MDK.csv', index_col=None, header=0, sep='\t')    #i.e. mesh disease key, needed?
MDM=pd.read_csv('MDM.csv',index_col=None, header=0, sep='\t') #mesh disease matches
doihere=MDM.iloc[:,2] #add doi.org/ check the name and match to tot_res
dh2=[]
for i in doihere.index:
    dh2.append('http://doi.org/'+doihere.loc[i])
from bs4 import BeautifulSoup
from bs4.dammit import EncodingDetector
import requests
#Frist we need a function:
soup=[]
for i in range(len(dh2)):
    parser = 'html.parser'  # or 'lxml' (preferred) or 'html5lib', if installed
    try:
        resp = requests.get(dh2[i])
        http_encoding = resp.encoding if 'charset' in resp.headers.get('content-type', '').lower() else None
        html_encoding = EncodingDetector.find_declared_encoding(resp.content, is_html=True)
        encoding = html_encoding or http_encoding
        soup.append(BeautifulSoup(resp.content, parser, from_encoding=encoding))
    except:
        soup.append('nan')
ohx=[]
for i in range(len(soup)):   
    ohx.append(soup[i].find('title'))
#    https://stackoverflow.com/questions/35956045/extract-title-with-beautifulsoup/35956388
ohion=[]
for i in range(len(ohx)):    
#    print(title) # Prints the tag, it could be this too..
#    ohio.append((oh[i].string)) # Prints the tag string content
#    txt = oh[i]
    if not isinstance(ohx[i], int):
        ohion.append(ohx[i].string.split("|"))
    else:
        ohion.append(ohx[i])
#%
import json
with open('test_tikka61020.txt', 'w') as f:
    f.write(json.dumps(ohion))
import json
with open('test_tikka23920a.txt', 'r') as f:
    ohiox = json.loads(f.read()) #work        
MDMsep3=[]
for i in range(len(ohiox)):
    if not isinstance(ohiox[i], int):
        MDMsep3.append(ohiox[i][0][:-1])
    else:
        MDMsep3.append(ohiox[i])   
#To select rows whose column value is in list 
#https://cmdlinetips.com/2018/02/how-to-subset-pandas-dataframe-based-on-values-of-a-column/
tot_res2.rename(columns={'Title of Article':'jname2'}, inplace=True)
tr_aha=tot_res2[tot_res2.jname2.isin(MDMsep3)]
okn=tr_aha['Review Word Count']<500    #vary this, float('Inf')
ok2=tr_aha.loc[okn]
blok=list(ok2['jname2']) #and then back to MDM        
#% This is how to get the indeces to the MDM ??
matsi=[]
satsi=[]
for i in range(len(MDMsep3)):
    for j in range(len(blok)):
        if MDMsep3[i]==blok[j]:
            matsi.append(i) #list(ok2.index)    
            satsi.append(MDMsep3[i])
MDM2=MDM.iloc[np.unique(matsi),:]              
dip=list(MDM2.index)
indexa=list(tuple(range(0,1428)))
matcha=set(dip) & set(indexa)  
M3=pd.DataFrame(MDMsep3)
M4=M3.loc[matcha,:]
MDM2['name'] = M4
#https://www.geeksforgeeks.org/adding-new-column-to-existing-dataframe-in-pandas/
testnm=MDM2.iloc[:,6:]
tak=testnm.iloc[:,:-1]
#% Our list:
#tot_res2.rename(columns={'Title of Article':'jname2'}, inplace=True)
#To select rows whose column value is in list 
#https://cmdlinetips.com/2018/02/how-to-subset-pandas-dataframe-based-on-values-of-a-column/
#tr_aha=tot_res2[tot_res2.jname2.isin(matsi)]
#okn=tr_aha['Review Word Count']<500
#ok2=tr_aha.loc[okn]
#https://stackoverflow.com/questions/6142689/initialising-an-array-of-fixed-size-in-python
#https://stackoverflow.com/questions/45473330/creating-a-pandas-data-frame-of-a-specific-size
#https://www.geeksforgeeks.org/adding-new-column-to-existing-dataframe-in-pandas/   
#%https://stackoverflow.com/questions/16327055/how-to-add-an-empty-column-to-a-dataframe
#for i in range(15,40):
#    ok2[i] = "" 
#%
ok2=ok2.reset_index()
testnm=testnm.reset_index()
pist=[]
for i in ok2.index:
    for j in testnm.index: 
        if ok2['jname2'].loc[i]==testnm['name'].loc[j]:
            pist.append([j])
ppist=[] 
for i in range(len(pist)):        
    ppist.append(pist[i][0])
from collections import Counter
z=Counter(ppist)
#https://www.w3schools.com/python/python_howto_remove_duplicates.asp
okj=[]
mylist = z
mylist = list(dict.fromkeys(mylist))
okj=mylist            
pix=[]
for i in range(len(z)):
    for j in range(0,len(testnm)):
        for n in range(0,):
            pix.append(testnm.iloc[j,:-1])
            pix.append(testnm.iloc[j,:-1])
indexia=[]
for i in range(0,len(testnm)):
    indexia.append([i] * int(z[okj[i]]))    
indexia2=[item for sublist in indexia for item in sublist]          
testaN=[]
for i in range(0,len(ok2)):
    testaN.append(testnm.loc[indexia2[i]])
testXX=pd.DataFrame(testaN)
testXX=testXX.reset_index()
ok2=ok2.reset_index()
ok2=ok2.iloc[:,:15]
testXY=testXX.iloc[:,2:-1]
lopi=[ok2.T,testXY.T]
topi=pd.concat(lopi)  
topi2=topi.T
#%yes, now I got it
#%Preparoi topi2.. ehkei tarpeen..
topi2=topi2.iloc[:,2:]
ina=topi2.iloc[:,13:].astype(str).astype(int)
topi22=topi2.iloc[:,:13]
nopi=[topi22.T,ina.T]
topi3=pd.concat(nopi) 
topi3=topi3.T
#%
def cat_word_mean(cat=13):
    avg1=[]
#    cat=13
    for i in range(0,len(topi2)): #len(topi2.iloc[0,13:])
        if topi2.iloc[i,cat]==1:
            avg1.append(topi2.loc[i,'Review Word Count'])
    avg2=[]
    for i in range(len(avg1)):
        avg2.append(float(avg1[i]))   
    #https://stackoverflow.com/questions/15033511/compute-a-confidence-interval-from-sample-data
    import numpy as np
    import scipy.stats
    def mean_confidence_interval(data, confidence=0.95):
        a = 1.0 * np.array(data)
        n = len(a)
        m, se = np.mean(a), scipy.stats.sem(a)
        h = se * scipy.stats.t.ppf((1 + confidence) / 2., n-1)
        return m, m-h, m+h   
    mci=mean_confidence_interval(avg1, confidence=0.95)
    return mci
#%
means_cat=[]    
for i in range(0,24):
    means_cat.append(cat_word_mean(cat=13+i))
loli=pd.DataFrame(means_cat)
#list(data.columns.values) 
#https://www.geeksforgeeks.org/python-change-column-names-and-row-indexes-in-pandas-dataframe/
data=topi2.iloc[0:3,13:]
loli.index=list(data.columns.values)
loli.columns =['Mean', 'CI_neg', 'CI_pos']
loli.to_csv('cat_word_meansmax_tikka291020.csv',index=True)
nti=list(topi2.loc[:,'jname2'])
txns=topi2
#https://www.geeksforgeeks.org/python-change-column-names-and-row-indexes-in-pandas-dataframe/
txns.index=nti
#%Jei hou!! :) ->
#prt=len(txns.loc[np.unique(txns.index)[0]])
#% You need the tnx dataframe for this:
def rev_counta(cat=13):  
    #%
    rev_count=[]
    rc_name=[]
#    cat=22
    for i in range(len(np.unique(txns.index))): #len(topi2.iloc[0,13:])
        if isinstance(txns.ix[np.unique(txns.index)[i],cat], int):
            #https://stackoverflow.com/questions/3501382/checking-whether-a-variable-is-an-integer-or-not
            if txns.ix[np.unique(txns.index)[i],cat]==1:  
                rev_count.append(1)
                rc_name.append(np.unique(txns.index)[i])
        elif len(txns.ix[np.unique(txns.index)[i],cat])>1:
            if pd.DataFrame(txns.ix[np.unique(txns.index)[i],cat]).iloc[0,0]==1:
                rev_count.append(len(txns.ix[np.unique(txns.index)[i],cat]))
                rc_name.append(np.unique(txns.index)[i])
    revcat=pd.DataFrame(rev_count)
    revcat.index=rc_name
#    import numpy as np
#    https://stackoverflow.com/questions/13355517/python-free-variable-numpy-referenced-before-assignment-in-enclosing-scope
    import scipy.stats
    def mean_confidence_interval(data, confidence=0.95):
        a = 1.0 * np.array(data)
        n = len(a)
        m, se = np.mean(a), scipy.stats.sem(a)
        h = se * scipy.stats.t.ppf((1 + confidence) / 2., n-1)
        return m, m-h, m+h   
    jaahas=mean_confidence_interval(revcat, confidence=0.95)
    if str(list(jaahas)[0])=='nan':
        aps=[]
        aps=[0, 0, 0]
    elif str(list(jaahas)[0])!='nan':
        aps=[]
        for i in range(0,3):
            aps.append(float(mean_confidence_interval(revcat, confidence=0.95)[i]))
    #%
    return aps, txns.ix[:,cat].name 
rev_cata=[]
rc_nam=[]    
for i in range(0,24):
    rev_cata.append(rev_counta(cat=13+i)[0])
    rc_nam.append(rev_counta(cat=13+i)[1])
japa=pd.DataFrame(rev_cata)
japa.index=rc_nam    
#https://www.geeksforgeeks.org/python-change-column-names-and-row-indexes-in-pandas-dataframe/
japa.columns =['Mean', 'CI_neg', 'CI_pos']
japa.to_csv('cat_rev_meansmax_tikka291020.csv',index=True) 


#%% Printing the small review files:
#Not sure if the below cut is needed(?)
small=tot_res2[tot_res2.loc[:,'Review Word Count']<50] 
#smallok=np.sort(small.loc[:,'Review Word Count']) 
#https://pandas.pydata.org/pandas-docs/stable/reference/api/pandas.DataFrame.sort_values.html
smallok=small.sort_values(by=['Review Word Count'])
#%'nan' values dissappear in excel so I replace them with NA
#https://datatofish.com/replace-nan-values-with-zeros/
smallok['Reviewer Name']=smallok['Reviewer Name'].replace(np.nan, 'NA') 
smallok['Reviewing Date']=smallok['Reviewing Date'].replace(np.nan, 'NA')
smallok["Reviewer's Title"]=smallok["Reviewer's Title"].replace(np.nan, 'NA')
smallok["Reviewer's Institution"]=smallok["Reviewer's Institution"].replace(np.nan, 'NA')
smallok['Link to PDF of Reviewer']=smallok['Link to PDF of Reviewer'].replace(np.nan, 'NA')
smallok['Page Count']=smallok['Page Count'].replace(np.nan, 'NA')
smallok.to_csv('smallok.csv',index=False)     
#%
import pandas as pd
f =  pd.read_csv('below_50.txt', header=None)
#https://stackoverflow.com/questions/21546739/load-data-from-txt-with-pandas
#https://www.quora.com/How-do-I-open-a-URL-in-Google-Chrome-in-new-tab-using-Python
import webbrowser
for i in range(len(f)):
    url = f.iloc[i][0]
    webbrowser.open_new_tab(url) 
